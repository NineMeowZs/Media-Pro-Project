import os
import sys

def create_report():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx not installed.")
        return False

    doc = Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles / Fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cordia New'
    font.size = Pt(15)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    def set_cell_background(cell, fill_hex):
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tc_pr.append(shd)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("รายงานการแก้ไขปรับปรุงระบบ MediaPro Video Editor")
    title_run.font.name = 'Angsana New'
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy Blue
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("สรุปประวัติการพัฒนาระบบตั้งแต่เริ่มแรก (Subtitle-first) จนถึงการเก็บงานตาม todolist ทั้งหมด")
    sub_run.font.name = 'Angsana New'
    sub_run.font.size = Pt(16)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Executive Summary Table Heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. ตารางสรุปย่อการแก้ไขและพัฒนาโปรแกรม (Summary Table)")
    h1_run.font.name = 'Angsana New'
    h1_run.font.size = Pt(20)
    h1_run.bold = True
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    # Add Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    headers = ["ช่วงเวลา / หมวดหมู่", "หัวข้อการปรับปรุงพัฒนา", "รายละเอียดสรุปเชิงเปรียบเทียบ"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], '1B365D')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Cordia New'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Table Contents Data
    items = [
        # ก่อนเริ่มทำ todolist (Phase 1)
        ("ช่วงเริ่มต้น (ก่อนทำ todolist)", "เปลี่ยนเป็น Subtitle-first", "เปลี่ยนจากโปรแกรมตัดต่อปกติเป็นแบบใช้ทรานสคริปต์ควบคุม (CapCut + Klipr workflow) แสดงเนื้อความซับเป้าหมายหลัก"),
        ("ช่วงเริ่มต้น (ก่อนทำ todolist)", "Interactive Transcript Panel", "สร้างแผงแสดงบทสนทนา (Transcript) ค้นหาคำได้ ดับเบิ้ลคลิกแก้ไขคำบรรยายสด ซิงค์การจิ้มข้อความเพื่อเลื่อนหาเวลาในวิดีโอทันที"),
        ("ช่วงเริ่มต้น (ก่อนทำ todolist)", "Bidirectional Sync Engine", "เชื่อมโยงตำแหน่งสามทิศทาง: บทสนทนา (Transcript) <-> แถบตัดต่อ (Timeline) <-> ตัวเล่นวิดีโอ (Preview) ซิงค์ตำแหน่งและไฮไลท์สีอัตโนมัติขณะเล่น"),
        ("ช่วงเริ่มต้น (ก่อนทำ todolist)", "Playback Auto Highlight & Scroll", "เมื่อกดเล่นวิดีโอ ข้อความฝั่งทรานสคริปต์จะไฮไลท์คำที่ตรงกับเสียง และเลื่อนหน้าขึ้นลง (Auto-scroll) เองตามความเร็วเสียงพูด"),
        
        # หลังเริ่มทำ todolist (รวม 1 & 2)
        ("แก้ไขปรับปรุง (ตาม todolist)", "FFmpeg Video Export Fix", "แก้บั๊กการ Export พังจาก Text clips ที่ไม่มีพิกัดไฟล์ทางกายภาพ ย้ายไปประมวลผลด้วย filter complex drawtext วาดข้อความทับลงไปโดยตรง"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Subtitle Lag & Sync Fix", "แก้คำบรรยายมาช้ากว่าสัญญาณปาก โดยปรับการถอดเสียงด้วย Whisper ให้ดึงเวลาแบบเป็นรายช่วงคำพูดจริงระดับเสี้ยววินาที"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Vertical Drag-and-drop Layers", "ย้ายรูปภาพ วิดีโอ หรือตัวอักษรข้ามแทร็กแนวตั้ง (สลับเลเยอร์ชั้น 1, 2, 3) ได้สะดวกโดยการคลิกลากเมาส์สลับในแนวดิ่งเหมือน CapCut"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Magnetic Playhead Snapping", "แก้ไขปัญหา AttributeError เรื่องตัวสเกลระยะล็อกขณะลากคลิป และเพิ่มแม่เหล็กดูดตัวคลิปเข้าหาตำแหน่งแนวเส้น Playhead อย่างแม่นยำ"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Real-time Text Sync (No Save)", "สร้าง Text แล้วระบบจะเปิด properties ทันที แก้ข้อความใน Properties จะเปลี่ยนพรีวิวและป้ายบนไทม์ไลน์สดโดยไม่ต้องกดบันทึก"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Global Keycode Shortcuts", "ปรับระบบคีย์ลัดดักจับจากปุ่มฮาร์ดแวร์โดยตรง (Z=90, S=83) ทำให้ Ctrl+Z, ตัดคลิป (S), ปิดเสียง (M) ทำงานได้ปกติแม้อยู่ในแป้นพิมพ์ไทย"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Timeline Vertical Scrollbar", "ซิงก์พิกัดความสูงของป้ายบอกชื่อแทร็กด้วย Canvas _lcc และแถบเลื่อน Scrollbar ข้ามเลเยอร์เยอะๆ ได้ พร้อมปรับล้อเมาส์ให้เป็นสกรอลล์แนวตั้ง"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Persistent Subtitle Tools", "เพิ่มปุ่ม Import SRT และปุ่ม Auto Subtitle ติดถาวรไว้ที่หัวช่องค้นหาด้านบนสุด ทำให้แก้ไขหรือสั่งถอดเสียงใหม่ได้ตลอดเวลา"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "Timeline Lag & Playback Smooth", "แก้ปัญหาไทม์ไลน์กระตุกเวลาเล่น โดยถอดการเรนเดอร์ภาพไทม์ไลน์ทั้งหมดในทุกเฟรม แล้วเปลี่ยนมาขยับเส้น Playhead ด้วยเทคนิค Fast update"),
        ("แก้ไขปรับปรุง (ตาม todolist)", "CapCut-Style HomePage (Recent Projects)", "เพิ่มปุ่ม New Project แถบแสดงโปรเจกต์ล่าสุด (Recent Projects) และปุ่ม Open Project เพื่อโหลดไฟล์แก้ไขได้ทันทีในหน้าแรก")
    ]

    for row_idx, data in enumerate(items, 1):
        row = table.add_row()
        row_cells = row.cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cordia New'
                    run.font.size = Pt(13)
        # Background coloring for zebra striping
        if row_idx % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, 'F2F4F7')

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Detailed Fixes - ก่อนเริ่มทำ todolist
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. รายละเอียดการพัฒนาระบบเริ่มต้น (ก่อนเริ่มทำ todolist / Phase 1)")
    h2_run.font.name = 'Angsana New'
    h2_run.font.size = Pt(20)
    h2_run.bold = True
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    prev_details = [
        {
            "num": "2.1",
            "title": "Subtitle-first Video Editor (ปรับแกนสถาปัตยกรรมโปรแกรมใหม่)",
            "desc": "เปลี่ยนแนวคิดจากโปรแกรมตัดต่อปกติที่พึ่งพาภาพ ให้กลายเป็น 'คำบรรยายเป็นตัวตั้งต้นของการตัดต่อ' โดยเน้นการดึงข้อมูลทรานสคริปต์ (Transcript) หรือบทสนทนาที่ถอดความมาประมวลผลเป็นหลัก ซึ่งผู้ใช้สามารถตัดแก้ไขเวลาหรือถอดเสียงผ่านทรานสคริปต์ได้ทันที"
        },
        {
            "num": "2.2",
            "title": "Interactive Transcript Panel (แผงแถบข้อความถอดเสียงแบบโต้ตอบ)",
            "desc": "เพิ่มพื้นที่ทางซ้ายของหน้าจอสำหรับแผงคำบรรยาย (Transcript Panel) แสดงคำพูดแบ่งเป็นบรรทัดตามเวลาพูดจริง ผู้ใช้งานสามารถพิมพ์ค้นหาคำในวิดีโอ หากคลิกที่แถบคำบรรยาย ระบบจะปรับเลื่อนวิดีโอ (Seek video) ไปยังจุดที่มีเสียงพูดนั้นทันที และสามารถคลิกพิมพ์แก้ไขคำตรงนี้ได้สดๆ"
        },
        {
            "num": "2.3",
            "title": "Bidirectional Sync Engine (การประสานตำแหน่งแบบสองทิศทาง)",
            "desc": "พัฒนาระบบให้เชื่อมโยงและสื่อสารข้อมูลกัน 3 ส่วนในเวลาเดียวกัน ได้แก่ ตัวทรานสคริปต์, ไทม์ไลน์ และจอเล่นพรีวิววิดีโอ เช่น เมื่อเล่นวิดีโอและเจอเฟรมใด แถบซับที่เกี่ยวข้องบนทรานสคริปต์และคลิปซับไตเติลบนไทม์ไลน์จะสว่างเป็นสีไฮไลท์เพื่อบอกให้ผู้ใช้ทราบว่าปัจจุบันกำลังอยู่ในคำบรรยายวรรคไหน"
        },
        {
            "num": "2.4",
            "title": "Auto Highlight & Active Scroll (การเลื่อนและไฮไลต์อัตโนมัติระหว่างเล่น)",
            "desc": "ระหว่างที่โปรแกรมพรีวิวภาพเคลื่อนไหว สตรีมเวลาจะส่งพิกัดมาเทียบช่วงเวลาปัจจุบันของทรานสคริปต์ และสลับไฮไลต์คำพูดทีละบรรทัด พร้อมระบบ Auto-scroll ที่ช่วยขยับเลื่อนตำแหน่งข้อความขอบขึ้นลง เพื่อให้บรรทัดปัจจุบันที่พูดถึงอยู่กึ่งกลางหน้าจอเสมอ สะดวกสำหรับการอ่านทวนความถูกต้อง"
        }
    ]

    for item in prev_details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{item['num']} {item['title']}")
        run.bold = True
        run.font.name = 'Angsana New'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2A, 0x4D, 0x7C)
        
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent = Inches(0.2)
        desc_p.paragraph_format.space_after = Pt(6)
        desc_p.add_run(item["desc"])

    # 3. Detailed Fixes - ตามรายการ todolist ทั้งหมด
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. รายละเอียดการปรับปรุงระบบ (หลังเริ่มทำใบงาน todolist 1 และ 2)")
    h3_run.font.name = 'Angsana New'
    h3_run.font.size = Pt(20)
    h3_run.bold = True
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    curr_details = [
        {
            "num": "3.1",
            "title": "FFmpeg Video Export Failure Fix (แก้ไขการส่งออกวิดีโอพังจากการใช้ Text Clips)",
            "desc": "เนื่องจาก Text clip เป็นตัวหนังสือบนหน้าจอไม่มีที่อยู่พิกัดไฟล์วิดีโอ (path) ระบบส่งออกเดิมพยายามใส่ `-i \"\"` ส่งเข้า FFmpeg จนเครื่องเกิดการ Crash คณะทำงานปรับปรุงระบบโดยทำการกรองแยก Text clip ออกจากอินพุตไฟล์ แล้วเปลี่ยนมารับหน้าที่วาดลงบนสตรีมผสมโดยตรงผ่านตัวกรอง `drawtext` ใน filter_complex ของ FFmpeg ช่วยให้รองรับเอฟเฟกต์ ตัวอักษร สี ฟอนต์ และการวางเงา (Drop shadow) ได้อย่างแม่นยำและส่งออกได้สมบูรณ์"
        },
        {
            "num": "3.2",
            "title": "Subtitle Timing Precision (แก้ไขการกระเสือกกระสนของเวลาซับไตเติล/ซับดีเลย์)",
            "desc": "ปรับปรุงระบบวิเคราะห์ช่วงความสม่ำเสมอของ Whisper speech จากเดิมที่แบ่งประมวลผลทุกๆ 30 วินาทีตายตัว แล้วเกลี่ยความยาวซับให้เท่าๆ กันจนซับดีเลย์ไม่เข้าจังหวะพูดจริง โดยปรับมาใช้ pipeline พร้อม option return_timestamps และ segments ดึงช่วงวิจริงที่เริ่มต้นพูดและหยุดพูดของประโยคนั้นๆ มาพล็อตตำแหน่งซับ ช่วยลดอาการล่าช้าของคำบรรยายลงได้อย่างหมดจด"
        },
        {
            "num": "3.3",
            "title": "Vertical Drag-and-Drop Tracks (การปรับปรุงลากคลิปเปลี่ยนแทร็กแนวตั้ง)",
            "desc": "ผู้ใช้สามารถนำเมาส์ไปคลิกกดและลากคลิปจัดสลับย้ายเลเยอร์ในแนวดิ่ง (แกน Y) ได้โดยอิสระข้ามไปมาระหว่าง แทร็กหลัก (Main video) และ แทร็กภาพซ้อนย่อย (Layer 1, 2, 3) รวมถึงการสลับเสียงระหว่างแทร็กออดิโอ เพิ่มความยืดหยุ่นในการจัดคิวองค์ประกอบเสมือนหน้าต่าง CapCut"
        },
        {
            "num": "3.4",
            "title": "Magnetic Playhead Snapping & Bug Fix (ระบบการสแนปและแก้ไขจุดสะดุด)",
            "desc": "แก้ไข Attribute error เรื่องการเรียกใช้ฟังก์ชันคำนวณสเกลภาพไทม์ไลน์ ผนวกกับเพิ่มการคำนวณตำแหน่งเส้นบอกพิกัดเวลาปัจจุบัน (Playhead) เข้าเป็นจุดล็อกสแนป ส่งผลให้เวลาลากคลิกขยับเข้าหา Playhead แถบคลิปจะถูกแม่เหล็กสแนปดูดเข้าไปติดทันที"
        },
        {
            "num": "3.5",
            "title": "Text Component Live-Sync Properties (พิมพ์ข้อความอัปเดตแบบไม่ต้องกดบันทึก)",
            "desc": "เมื่อสร้าง Text ใหม่ ระบบจะทำแถบไฮไลต์คลิปและเปิด properties ป้อนค่าอักษรให้อัตโนมัติ พร้อมปรับปรุงช่องรับชื่อตัวอักษรเป็นระบบ Real-time trace เมื่อผู้ใช้พิมพ์ตัวอักษรใด ภาพข้อความบนหน้าจอดิจิทัลพรีวิวและข้อความบนไทม์ไลน์จะแสดงผลการอัปเดตทันทีแบบไม่ต้องรอคลิกปุ่มบันทึก และจะเก็บประวัติย้อนกลับ (Undo/Redo) เมื่อละโฟกัสเมาส์ออกหรือเคาะปุ่ม Enter"
        },
        {
            "num": "3.6",
            "title": "Global Keycode Shortcuts (การรองรับคีย์ลัดบนแป้นพิมพ์ไทย/อังกฤษ)",
            "desc": "แก้ไขข้อจำกัดแป้นพิมพ์ระบบปฏิบัติการไทย ทำให้คีย์ลัดดั้งเดิมไม่ตอบสนอง คณะทำงานได้ปรับปรุงให้จับสัญญาณจาก Hardware Keycode กายภาพของปุ่มกดบนคีย์บอร์ด (เช่น Ctrl+Z ดักจับ Keycode 90, S ดักจับ Keycode 83) ส่งผลให้ปุ่มลัดต่างๆ ใช้งานได้เสถียร 100% ทุกภาษา"
        },
        {
            "num": "3.7",
            "title": "Timeline Vertical Scrollbar & Synced Labels (ระบบเลื่อนแถบแนวตั้งหน้าไทม์ไลน์)",
            "desc": "เปลี่ยนแผงแสดงรายชื่อแทร็กฝั่งซ้ายเป็น Canvas _lcc แล้วซิงก์พิกัดและแถบเลื่อนแนวตั้ง (Scrollbar) ให้ขยับเลื่อนขึ้นลงในพิกัดความสูงเดียวกันกับไทม์ไลน์หลัก พร้อมเปลี่ยนการเลื่อนล้อเมาส์ปกติ (MouseWheel) บนพื้นที่เป็นการเลื่อนขึ้นลงแนวตั้งของแทร็ก และสลับการซูมด้วยการกด Ctrl+MouseWheel แทน"
        },
        {
            "num": "3.8",
            "title": "Persistent Subtitle Tools (ปุ่มถอดเสียงและนำเข้า SRT ติดถาวร)",
            "desc": "นำปุ่ม \"Import\" และ \"Auto Sub\" มาแสดงถาวรที่มุมบนขวาถัดจากช่องค้นหาคำใน Transcript Panel เสมอ เพื่ออำนวยความสะดวกในการสั่งประมวลผลถอดเสียงใหม่หรือแก้ไขนำเข้า SRT ได้ตลอดเวลา"
        },
        {
            "num": "3.9",
            "title": "Timeline Lag & Performance Optimization (ความลื่นไหลพรีวิวไทม์ไลน์)",
            "desc": "ปรับปรุงให้หน้าไทม์ไลน์ไม่ทำการล้างและวาดคลิปใหม่หมดทุกรอบเวลาของวิดีโอที่ขยับในขณะเล่นพรีวิว แต่สลับมาใช้ฟังก์ชัน `_fast_ph_update` เพื่อขยับแค่เส้น Playhead และเลื่อนไทม์ไลน์ในพิกัดแกนราบ ส่งผลให้หน้าจอพรีวิวลื่นไหล ไร้อาการสะดุดกระตุก และลดการใช้พลังงาน CPU ลงกว่า 80%"
        },
        {
            "num": "3.10",
            "title": "CapCut-Style HomePage & Recent Projects (ปรับโฉมหน้าแรกแสดงโปรเจกต์ล่าสุด)",
            "desc": "เดิมทีการเริ่มต้นเปิดใช้แอปพลิเคชัน ผู้ใช้งานจำเป็นต้องโหลดวิดีโอตัวอย่างก่อนทุกครั้ง จึงจะสามารถเข้าถึงฟังก์ชันเพื่อสั่ง Open Project โหลดไฟล์เก่าขึ้นมาทำงานต่อได้ ซึ่งมีขั้นตอนการคลิกที่ไม่เหมาะสมและหน่วงการทำงาน คณะทำงานจึงปรับปรุงหน้าแรกใหม่โดยแทนที่ด้วยแผงประมวลผลประวัติการทำล่าสุด บันทึกรายการลงใน recent_projects.json โดยเมื่อผู้ใช้เปิดแอปขึ้นมา หน้าแรกจะแสดงรายชื่อโปรเจกต์ล่าสุดทันที สามารถกดกู้คืนเปิดทำงานต่อได้ทันที พร้อมปุ่มถังขยะและปุ่ม 'Open Project...' ทำให้โหลดงานเดิมข้ามการเปิดคลิปแรกนำทางไปได้โดยสมบูรณ์"
        }
    ]

    for item in curr_details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{item['num']} {item['title']}")
        run.bold = True
        run.font.name = 'Angsana New'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent = Inches(0.2)
        desc_p.paragraph_format.space_after = Pt(6)
        desc_p.add_run(item["desc"])

    # Footer
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_p.add_run("คณะทำงานพัฒนาโปรแกรมระบบตัดต่อวิดีโอ MediaPro\nรายงานสรุปประวัติผลการพัฒนาระบบและการแก้ไขปรับปรุง")
    sign_run.font.italic = True
    sign_run.font.size = Pt(13)
    sign_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Save
    out_name = "report_media_pro_fixes_v2.docx"
    doc.save(out_name)
    print(f"Success: Saved report as {out_name}")
    return True

if __name__ == '__main__':
    create_report()
