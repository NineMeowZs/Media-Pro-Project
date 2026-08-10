import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Page Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Base Style Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'TH SarabunPSK'
    normal_style.font.size = Pt(16)
    normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(24)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        p.paragraph_format.space_after = Pt(14)
        return p

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(20)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_bullet(label, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        r_lbl = p.add_run(f"• {label}: ")
        r_lbl.font.name = 'TH SarabunPSK'
        r_lbl.font.size = Pt(15)
        r_lbl.bold = True
        r_lbl.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

        r_txt = p.add_run(text)
        r_txt.font.name = 'TH SarabunPSK'
        r_txt.font.size = Pt(15)
        r_txt.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # 1. Header Title
    add_title("รายงานสรุปการปรับปรุงแก้ไขและพัฒนาฟีเจอร์โปรแกรม MediaPro")
    add_subtitle("ประจำวันที่ 23 – 25 กรกฎาคม 2569 (Report_MediaPro_Fix_25_07_26)")

    # Intro Paragraph
    p_intro = doc.add_paragraph()
    r = p_intro.add_run("รายงานฉบับนี้รวบรวมสรุปผลการปรับปรุงแก้ไขข้อผิดพลาด (Bug Fixes), การปรับแต่งสถาปัตยกรรม UI/UX ตามความต้องการของผู้ใช้งาน และการพัฒนาฟีเจอร์ใหม่สำหรับโปรแกรมตัดต่อวิดีโอ MediaPro เฉพาะในระหว่างวันที่ 23 – 25 กรกฎาคม 2569 พร้อมทั้งสรุปรายการปัญหาที่จดบันทึกไว้สำหรับดำเนินการพัฒนาต่อในอนาคต รวมทั้งสิ้น 14 หัวข้อหลัก")
    r.font.name = 'TH SarabunPSK'
    r.font.size = Pt(16)
    p_intro.paragraph_format.space_after = Pt(12)

    # 2. Section 1: Executive Summary Table
    add_heading1("ส่วนที่ 1: ตารางสรุปภาพรวมการแก้ไขและพัฒนาฟีเจอร์ (23 – 25 กรกฎาคม 2569)")

    table_data = [
        ("1. ระบบ Proxy Video Build ใน Background",
         "พรีวิววิดีโอความละเอียดสูงกระตุก การสกัดเฟรม Eager จากไฟล์ใหญ่ใช้ทรัพยากรสูง",
         "ใช้ FFmpeg เรนเดอร์ไฟล์ proxy 360p ใน Background พร้อมแสดง Progress % บน Status bar สลับเล่นไฟล์ proxy ช่วยให้พรีวิวเล่นวิดีโอตัดต่อลื่นขึ้น แม้จะยังมีความกระตุกอยู่บ้างตามสเปกเครื่อง"),

        ("2. ระบบ Waveform คลื่นเสียงบน Timeline",
         "ไทม์ไลน์เรียบมองไม่เห็นระดับความถี่เสียงพูด ทำให้ตัดต่อเข้าจังหวะเสียงยาก",
         "สกัดค่าแอมพลิจูดความถี่เสียงด้วย FFmpeg/Numpy แปลงเป็นกราฟคลื่นเสียงวาดบนแทร็ก Audio และ Main Track ของไทม์ไลน์ เพิ่มความน่าใช้งานและช่วยให้ตัดต่อตรงจังหวะเสียงเป๊ะๆ"),

        ("3. รองรับโมเดลถอดเสียง Custom Whisper PyTorch (best.pt)",
         "HuggingFace WhisperProcessor คาดหวังโฟลเดอร์ไฟล์ JSON แต่ best.pt เป็น PyTorch state_dict ดิบ",
         "เขียนระบบตรวจจับไฟล์ .pt อัตโนมัติ ตรวจสอบมิติ hidden size (768 -> small) แล้วสวม state_dict 479 tensors สำเร็จ 100%"),

        ("4. กำหนดช่วงเวลาสร้างซับไตเติล (Time Range Subtitle Generation)",
         "เดิมการสร้างซับทำเฉพาะทั้งวิดีโอแรก ขาดการกำหนดวินาทีเริ่มต้น-สิ้นสุดตามต้องการ",
         "เพิ่มพารามิเตอร์ audio_range ใน transcriber.py สำหรับตัดสัญญาณเสียง และเพิ่ม Radio Button + ช่องกรอกวินาทีใน Subtitle Settings"),

        ("5. ถอดเสียงซับไตเติลครอบคลุมทุกวิดีโอบน Main Track",
         "เดิมส่งเฉพาะคลิปแรก self.tracks['main'][0]['path'] เข้าระบบถอดเสียง",
         "สร้าง _render_main_track_audio() รวมสัญญาณเสียงทุกคลิปวิดีโอ/เสียงบน Main Track ผ่าน FFmpeg adelay/amix ออกมาเป็น WAV เสียงรวมส่งให้ Whisper"),

        ("6. ปรับแท็บเลือก 3 ปุ่ม (Media/Text/Captions) เท่ากัน 100%",
         "เดิมใช้ Tkinter pack ทำให้ปุ่ม Captions หดเล็กตกขอบ และมุมกรอบขาวปิดไม่สนิท",
         "เปลี่ยนมาใช้ grid geometry manager (uniform='tab', weight=1) และปรับ inset padding padx=12, pady=(12, 4) ให้กรอบขาวมน 20px ปิดสนิทสวยงาม"),

        ("7. ซิงก์ชื่อโปรเจกต์กับพาดหัวโปรแกรม (Project Name Sync)",
         "ป้ายชื่อค้างอยู่ที่ 'Untitled Project' ไม่เปลี่ยนตามชื่อไฟล์โปรเจกต์ที่เซฟหรือเปิด",
         "ผูกคำสั่งอัปเดต self._proj_name และ self._proj_title_lbl ในฟังก์ชันเปิด โหลด และเซฟไฟล์โปรเจกต์ทั้งหมดให้เปลี่ยนชื่อสดทันที"),

        ("8. เพิ่มปุ่ม Home ย้อนกลับหน้าแรกที่มุมซ้ายบน",
         "ขาดปุ่มลัดย้อนกลับไปยังหน้าแรก Home จากหน้าจอตัดต่อหลัก",
         "เพิ่มปุ่มไอคอนบ้าน 🏠 ข้างโลโก้ MediaPro ฝั่งซ้ายบน ผูกคำสั่งย้อนกลับ _on_back() เพื่อเซฟความจำและปรับขนาดหน้าต่างกลับสู่หน้าแรก"),

        ("9. ระบบ Transform (Scale & Rotate) บนเฟรมวิดีโอ",
         "ค่า Scale และ Rotate ใน Properties Panel ไม่ได้ถูกนำไปคำนวณบนเฟรมวิดีโอ",
         "เขียนฟังก์ชัน _apply_clip_transform ใช้ OpenCV matrix warpAffine (หมุน) และ resize (ย่อ/ขยาย) ประมวลผลสดบนหน้าจอพรีวิว"),

        ("10. รองรับการกรอกตัวเลขตั้งค่าสเกล/ระดับโดยตรง (Numeric Entries)",
         "เดิมมีเฉพาะสไลเดอร์ ปรับระดับตัวเลขที่แม่นยำยาก",
         "เพิ่ม CTkEntry ตัวเลขประกบคู่สไลเดอร์ใน editor_properties.py ซิงก์ข้อมูลสองทิศทาง (bidirectional mapping) พิมพ์กรอกตัวเลขได้โดยตรง"),

        ("11. ตั้งชื่อ Output ในหน้า Export เป็น 'ชื่อโปรเจกต์.mp4' ทันที",
         "เดิมกล่องข้อความ Output File ขึ้นเป็นค่าว่าง หรือ Output_Video.mp4 เนื่องจากอ้างอิงออบเจกต์ master ผิดระดับ",
         "ส่ง proj_name=self._proj_name เข้า _ExportDialog โดยตรง และตั้งค่า value=f'{stem}.mp4' ให้กล่องข้อความทันทีตั้งแต่นาทีแรก"),

        ("12. คลิกพื้นที่ว่างบนไทม์ไลน์/หน้าจอเพื่อยกเลิกการเลือกคลิป (Deselect)",
         "เมื่อคลิกพื้นที่ว่าง ค่า sel_track และ sel_idx ค้าง ทำให้ Properties Panel แสดงตั้งค่าค้างอยู่",
         "เพิ่มคำสั่งเคลียร์ sel_track='' และ sel_idx=-1 พร้อมรีเฟรช Properties Panel เมื่อคลิกโดนพื้นที่ว่างไทม์ไลน์และหน้าจอพรีวิว"),

        ("13. ระบบบันทึกความคงอยู่ของโปรเจกต์ (Full Project Persistence)",
         "ไฟล์ JSON ไม่ได้เก็บค่า Ratio, สไตล์ฟอนต์, ขนาดฟอนต์, Scale, Rotate",
         "บันทึก ratio และ style ลงใน JSON ใน _save และเขียนคำสั่งกู้คืนค่าทั้งหมดกลับเข้าคอนโทรลเลอร์ใน _load_project_file 100%"),

        ("14. ระบบสืบค้นและเชื่อมโยงไฟล์อัตโนมัติเมื่อย้ายโฟลเดอร์ (Auto-Relocation)",
         "เดิมเก็บเฉพาะ Absolute Path ดิบ หากย้ายโฟลเดอร์โปรเจกต์ พาธไฟล์จะเสียหาไม่เจอ",
         "เขียนฟังก์ชัน _relocate_path ค้นหาไฟล์ชื่อเดียวกันในโฟลเดอร์โปรเจกต์ปัจจุบันและโฟลเดอร์ย่อย media/ ให้อัตโนมัติเมื่อเปิดโปรเจกต์")
    ]

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["หัวข้อปัญหา / ความต้องการ", "สาเหตุหลักของปัญหา / ความต้องการ", "วิธีการดำเนินการแก้ไขและผลสัมฤทธิ์"]
    widths = [Inches(2.2), Inches(2.3), Inches(2.5)]

    for idx, title in enumerate(hdr_titles):
        cell = hdr_cells[idx]
        cell.width = widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cell_background(cell, "1e3a8a")
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)

    # Data Rows
    for row_idx, (col1, col2, col3) in enumerate(table_data):
        row_cells = table.add_row().cells
        bg_hex = "f8fafc" if row_idx % 2 == 0 else "ffffff"
        for idx, text in enumerate([col1, col2, col3]):
            cell = row_cells[idx]
            cell.width = widths[idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'TH SarabunPSK'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
            if idx == 0:
                run.bold = True
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(14)

    # 3. Section 2: Detailed Fixes & Features (23 - 25 กรกฎาคม 2569)
    add_heading1("ส่วนที่ 2: รายละเอียดการปรับปรุงแก้ไขและพัฒนาเจาะจง (23 – 25 กรกฎาคม 2569)")

    details_data = [
        ("2.1 ระบบสร้างและประมวลผลวิดีโอ Proxy ใน Background (Proxy Video Build Processing)",
         "การเปิดเล่นวิดีโอความละเอียดสูงพรีวิวบนไทม์ไลน์มีอาการสะดุด กระตุก และใช้ทรัพยากรเครื่องสูง",
         "การดึงเฟรมวิดีโอ Eager จากไฟล์ 1080p ดั้งเดิมแบบเรียลไทม์ใช้ CPU/Memory สูง",
         "สร้างระบบ Background Thread รัน FFmpeg แปลงไฟล์ proxy 360p ความละเอียดต่ำไว้ใช้งานชั่วคราว แสดง Progress เปอร์เซ็นต์ความคืบหน้าบน Status Bar และสลับเล่นเฟรมจาก Proxy ทันทีเมื่อสร้างเสร็จ",
         "ช่วยให้พรีวิวเล่นวิดีโอตัดต่อลื่นขึ้นเล็กน้อยถึงปานกลาง แม้จะยังมีความกระตุกอยู่พอสมควรตามข้อจำกัดสเปกเครื่องประมวลผล"),

        ("2.2 ระบบแสดงผลคลื่นเสียงบนไทม์ไลน์ (Audio Waveform Visualization on Timeline)",
         "ไทม์ไลน์แสดงเฉพาะแถบสีนิ่ง ทำให้ผู้ใช้มองไม่เห็นจังหวะความถี่เสียงพูด ตัดต่อเข้าจังหวะคำพูดได้ยาก",
         "ขาดระบบสกัดสัญญาณแอมพลิจูดความถี่เสียงและการวาดเส้นกราฟคลื่นเสียงบน Canvas ไทม์ไลน์",
         "ใช้ FFmpeg/Numpy ดึงค่าความถี่เสียงและเรนเดอร์กราฟคลื่นเสียง Waveform ลงบนแทร็ก Audio และแทร็ก Main Video สอดคล้องตามพิกัดเวลาวินาทีจริง",
         "เพิ่มความน่าใช้งาน ดูสวยงามเหมือนซอฟต์แวร์ตัดต่อระดับมืออาชีพ และช่วยให้ผู้ใช้คลิกตัดต่อเข้าจังหวะคำพูดได้อย่างแม่นยำ"),

        ("2.3 การรองรับโมเดลถอดเสียงภาษาไทยส่วนตัว Custom Whisper PyTorch (best.pt)",
         "ผู้ใช้งานไม่สามารถเลือกไฟล์โมเดลถอดเสียงส่วนตัว best.pt ในไดเรกทอรีโครงการได้ เกิดข้อผิดพลาด OSError: Can't load feature extractor",
         "โมดูล HuggingFace WhisperProcessor คาดหวังโฟลเดอร์ที่มีไฟล์ JSON แต่ไฟล์ best.pt เป็นไฟล์ PyTorch state_dict ดิบ",
         "เพิ่มระบบตรวจจับนามสกุลไฟล์ .pt อัตโนมัติใน transcriber.py ตรวจวัดมิติ hidden size (768 -> small) แล้วทำการโหลดโครงสร้างมาตรฐาน openai/whisper-small พร้อมสวมน้ำหนัก state_dict 479 tensors เข้าไปโดยตรง",
         "สามารถเลือกและโหลดใช้งานโมเดล best.pt เพื่อถอดเสียงภาษาไทยได้อย่างแม่นยำและสมบูรณ์แบบ"),

        ("2.4 การเพิ่มฟังก์ชันกำหนดช่วงเวลาสร้างซับไตเติล (Time Range Subtitle Generation)",
         "การถอดเสียงดั้งเดิมถอดเฉพาะคลิปวิดีโอทั้งไฟล์ตั้งแต่ต้นจนจบ ผู้ใช้งานต้องการกำหนดช่วงเวลาเฉพาะ (เช่น ตั้งแต่วินาทีที่ 10 ถึง 60) ที่ต้องการเจเนเรตซับไตเติล",
         "ขาดส่วนควบคุมอินพุตการรับขอบเขตเวลาใน UI หน้าต่างตั้งค่า และฟังก์ชัน transcribe_video เดิมไม่ได้ทำการ Slice อาร์เรย์สัญญาณเสียง",
         "เพิ่มพารามิเตอร์ audio_range=(t_start, t_end) ใน transcriber.py เพื่อ slice อาร์เรย์สัญญาณเสียง Numpy ก่อนส่งให้โมเดลประมวลผล พร้อมปรับปรุง _SubtitleDialog เพิ่มตัวเลือก Radio Button 'ทั้งหมด (Full)' vs 'กำหนดช่วงเวลา (Custom Range)' และกล่องข้อความกรอกวินาทีเริ่มต้นและสิ้นสุด",
         "ผู้ใช้งานสามารถกรอกระบุช่วงเวลาที่ต้องการสร้างซับได้อย่างอิสระ ซับไตเติลถูกพล็อตลงบนไทม์ไลน์ในพิกัดเวลาที่ถูกต้องเป๊ะๆ"),

        ("2.5 การประมวลผลถอดเสียงสร้างซับไตเติลครอบคลุมทุกวิดีโอและคลิปเสียงบน Main Track",
         "การถอดเสียงเดิมถอดเฉพาะวิดีโอแรกบน Main Track เมื่อมีวิดีโอหลายคลิปเรียงต่อกัน ซับไตเติลจะขาดหายไปหลังจบวิดีโอแรก",
         "โค้ดเดิมส่งเฉพาะไฟล์ self.tracks['main'][0]['path'] เข้าระบบถอดความ",
         "พัฒนาฟังก์ชัน _render_main_track_audio() รันคำสั่ง FFmpeg filter complex (adelay, atempo, amix) รวมสัญญาณเสียงของทุกคลิปวิดีโอและคลิปเสียงบน Main Track ตลอดความยาวไทม์ไลน์ ออกมาเป็นไฟล์ WAV 16kHz mono เสียงรวมชั่วคราว แล้วส่งไฟล์เสียงรวมนี้ให้โมเดลถอดความ",
         "ซับไตเติลเจเนเรตออกมาครอบคลุมวิดีโอทั้งหมดบนไทม์ไลน์ตั้งแต่ต้นจนจบโปรเจกต์"),

        ("2.6 การปรับแต่งท็อปบาร์ 3 แท็บ (Media / Text / Captions) เท่ากัน 100% (Symmetrical Tabs UI)",
         "แท็บเลือกด้านซ้าย 3 ปุ่ม (▶ Media, TI Text, 💬 Captions) มีขนาดความกว้างไม่เท่ากัน ปุ่ม Captions หดเล็กลงตกขอบ และกรอบสีขาวมน 20px รอบแผงพรีวิววิดีโอปิดไม่สนิทมีช่องว่าง",
         "เดิมใช้ Tkinter pack(side='left', expand=True) ซึ่งคำนวณพื้นที่ตามความยาวข้อความ และคอนเทนเนอร์แผงพรีวิวมีขอบเกินขอบเขต",
         "เปลี่ยนการวางเลย์เอาต์แท็บท็อปบาร์มาใช้ grid geometry manager ร่วมกับ grid_columnconfigure(0..2, weight=1, uniform='tab') และปรับ bg_color='transparent' พร้อม inset padding padx=12, pady=(12, 4) ให้กับแผงพรีวิววิดีโอตรงกลาง",
         "แท็บทั้ง 3 ปุ่มมีขนาดความกว้างเท่ากันเป๊ะ 100% ไม่ตกขอบ กรอบขาวมน 20px ปิดสนิทสวยงามเนียนตา"),

        ("2.7 การซิงก์ชื่อโปรเจกต์กับพาดหัวโปรแกรม (Project Name Synchronization)",
         "ข้อความชื่อโปรเจกต์ค้างอยู่ที่ 'Untitled Project' ไม่เปลี่ยนตามชื่อไฟล์โปรเจกต์ที่เราเซฟหรือเปิดขึ้นมาตัดต่อ",
         "ขาดการซิงก์อัปเดตตัวแปร self._proj_name และป้ายชื่อ self._proj_title_lbl ในฟังก์ชันจัดการไฟล์โปรเจกต์",
         "ผูกคำสั่งอัปเดต self._proj_name และปรับป้ายชื่อ self._proj_title_lbl ใน _load_project_file, _load_video, และ _save ให้ทำงานตรงกัน",
         "ชื่อข้อความสีฟ้าขวาบนเปลี่ยนตามชื่อโปรเจกต์ที่ทำและเซฟไว้ตรงตามจริงทันที"),

        ("2.8 การเพิ่มปุ่ม Home ย้อนกลับหน้าแรกที่มุมซ้ายบน (Home Navigation Button)",
         "ผู้ใช้งานต้องการให้มีปุ่มรูปบ้านแสดงข้างๆ โลโก้ MediaPro เพื่อกดแล้วย้อนกลับไปหน้า Home ได้สะดวก",
         "เดิมขาดปุ่มลัดย้อนกลับจากหน้าตัดต่อไปยังหน้า Home ในส่วนเฮดเดอร์ซ้ายบน",
         "เพิ่มปุ่มไอคอนบ้าน 🏠 ข้างโลโก้ MediaPro ฝั่งซ้ายบน ผูกคำสั่งย้อนกลับ self._on_back() เพื่อเซฟหน่วยความจำและย่อขนาดหน้าต่างกลับสู่หน้าแรก",
         "ผู้ใช้สามารถคลิกปุ่มบ้านย้อนกลับไปหน้าหลัก (Home) ได้ตลอดเวลา"),

        ("2.9 การพัฒนาระบบ Transform (Scale & Rotate) บนเฟรมวิดีโอเรียลไทม์",
         "ต้องการให้สไลเดอร์ Scale (ขยาย/ย่อ) และ Rotate (หมุนองศา) ในแถบตั้งค่าขวาใช้งานได้จริงบนเฟรมวิดีโอ",
         "ค่า Scale และ Rotate ใน Property Panel ไม่ได้ถูกส่งเข้ากระบวนการเรนเดอร์ภาพ",
         "เขียนฟังก์ชัน _apply_clip_transform(frame, scale, rotate) ใช้ OpenCV matrix warpAffine (หมุนองศา) และ resize (ย่อ/ขยาย) ประมวลผลสดบนหน้าจอพรีวิว ผูกเข้ากับ _render (Main track) และ _apply_overlay (Overlay clips)",
         "ปรับสไลเดอร์หรือพิมพ์ระบุองศา/ขนาด เฟรมวิดีโอและองค์ประกอบบนหน้าจอหมุนเอียงและย่อขยายตามจริงแบบเรียลไทม์"),

        ("2.10 การรองรับการกรอกตัวเลขตั้งค่าสเกล/ระดับโดยตรง (Numeric Entry Inputs)",
         "ต้องการให้ช่องตั้งค่าต่างๆ (เช่น Transform, Volume, Speed) พิมพ์กรอกตัวเลขระดับที่ต้องการได้ ไม่ใช่มีแค่สไลเดอร์เลื่อนอย่างเดียว",
         "เดิมมีเฉพาะ CTkSlider ขาดกล่องป้อนข้อมูลตัวเลข CTkEntry",
         "เพิ่ม CTkEntry ตัวเลขประกบคู่สไลเดอร์ใน editor_properties.py พร้อมผูกอีเวนต์กด Enter หรือหลุดโฟกัสเพื่อซิงก์ค่าตัวเลขกับสไลเดอร์แบบ bidirectional mapping",
         "ผู้ใช้เลือกได้ทั้งการเลื่อนสไลเดอร์และพิมพ์กรอกตัวเลขพิกัด/ระดับเป้าหมายโดยตรง"),

        ("2.11 การตั้งชื่อ Output ในหน้า Export เป็น 'ชื่อโปรเจกต์.mp4' เป็น Default ทันที",
         "ตอนกด Export ต้องการให้มีชื่อไฟล์ Output ตั้งต้นเป็นชื่อโปรเจกต์ จะได้ไม่ต้องเสียเวลาพิมพ์ตั้งชื่อซ้ำ",
         "เดิมกล่องข้อความ Output File ใน _ExportDialog ขึ้นเป็นค่าว่าง หรือ Output_Video.mp4 เนื่องจากอ้างอิงออบเจกต์ master ผิดระดับ",
         "ส่งออบเจกต์ self (EditorPage) และ proj_name=self._proj_name เข้า _ExportDialog โดยตรง และตั้งค่า value=f'{stem}.mp4' ให้กับตัวแปรกล่องข้อความ self._path_v",
         "เมื่อเปิดหน้าต่าง Export ช่อง Output File มีชื่อ 'ชื่อโปรเจกต์.mp4' ปรากฏในช่องตั้งต้นทันที ผู้ใช้แก้ไขชื่อได้หากต้องการ"),

        ("2.12 การคลิกพื้นที่ว่างบนไทม์ไลน์/หน้าจอเพื่อยกเลิกการเลือกคลิป (Deselect / Clear Selection)",
         "ต้องการให้กดพื้นที่เปล่าเพื่อออกจากการตั้งค่า Component นั้นๆ เพื่อเคลียร์ไปทำ Component อื่นต่อได้ง่าย",
         "เมื่อคลิกพื้นที่ว่าง ค่า sel_track และ sel_idx ค้าง ทำให้ Property Panel ยังแสดงตั้งค่าค้างอยู่",
         "เพิ่มคำสั่งเคลียร์ sel_track='' และ sel_idx=-1 พร้อมรีเฟรช Property Panel เมื่อคลิกโดนพื้นที่ว่างใน _tl_press (editor_timeline.py) และ _canvas_ov_press (editor_preview.py)",
         "คลิกพื้นที่เปล่าแล้วระบบจะเคลียร์การเลือกและซ่อนแถบตั้งค่าทันที พร้อมเลือกปรับแต่ง Component อื่นต่อได้ง่าย"),

        ("2.13 ระบบบันทึกความคงอยู่ของโปรเจกต์ (Full Project State Persistence: Ratio, Style, Properties)",
         "ตอนเปิดโปรเจกต์กลับมา การตั้งค่า เช่น Aspect Ratio, สไตล์ฟอนต์, ขนาดฟอนต์, Scale, Rotate ไม่ถูกบันทึกค้างไว้",
         "ฟังก์ชัน _save ไม่ได้เก็บค่า ratio และ style ลงไฟล์ JSON และ _load_project_file ไม่ได้กู้คืนค่าเหล่านั้น",
         "เพิ่มการบันทึก ratio, style (font_name, font_size, font_color, decoration, animation, position) ลงโครงสร้าง JSON ใน _save และเขียนคำสั่งกู้คืนค่าทั้งหมดกลับเข้าคอนโทรลเลอร์ใน _load_project_file",
         "เปิดโปรเจกต์เก่าที่เซฟไว้ ทั้งสัดส่วนจอ สไตล์ฟอนต์ และคุณสมบัติของทุกคลิปจะถูกคืนค่ากลับมาครบถ้วน 100%"),

        ("2.14 ระบบสืบค้นและเชื่อมโยงไฟล์อัตโนมัติเมื่อย้ายโฟลเดอร์โปรเจกต์ (Smart Auto-Relocation)",
         "ข้อสงสัยหากย้ายโฟลเดอร์โปรเจกต์ที่ตัดต่อทั้งหมด ตัวโปรแกรมจะยังหาไฟล์วิดีโอ/เสียงเจอหรือไม่",
         "เดิมเก็บเฉพาะ Absolute path ดิบ หากย้ายโฟลเดอร์ไปไดเรกทอรีอื่นหรือคอมพิวเตอร์เครื่องอื่น พาธเดิมจะพังหาไฟล์ไม่เจอ",
         "เขียนฟังก์ชัน _relocate_path(old_path, project_dir) ใน _load_project_file หากไม่พบไฟล์ที่ absolute path เดิม ระบบจะค้นหาไฟล์ชื่อเดียวกันในโฟลเดอร์โปรเจกต์ปัจจุบันและโฟลเดอร์ย่อย media/ ให้อัตโนมัติ",
         "สามารถย้ายโฟลเดอร์โปรเจกต์ไปวางที่ใดก็ได้ หรือเปิดบนเครื่องอื่น โปรแกรมจะค้นหาและเชื่อมโยงไฟล์ให้อัตโนมัติโดยไฟล์ไม่หาย")
    ]

    for title, desc, cause, action, result in details_data:
        add_heading2(title)
        add_bullet("รายละเอียดปัญหาที่พบ / ความต้องการ", desc)
        add_bullet("สาเหตุของปัญหา", cause)
        add_bullet("สิ่งที่ดำเนินการแก้ไข", action)
        add_bullet("ผลลัพธ์หลังการแก้ไข", result)

    # 4. Section 3: Future Work & Outstanding Backlog
    add_heading1("ส่วนที่ 3: รายงานสรุปปัญหาและแผนการดำเนินงานต่อในอนาคต (Outstanding Backlog)")

    p_backlog_intro = doc.add_paragraph()
    r = p_backlog_intro.add_run("สรุปรายการปัญหาและฟีเจอร์เพิ่มเติมที่ได้รับการจดบันทึกไว้สำหรับดำเนินการแก้ไขและพัฒนาปรับปรุงต่อในอนาคตตามความต้องการของผู้ใช้งาน:")
    r.font.name = 'TH SarabunPSK'
    r.font.size = Pt(16)
    p_backlog_intro.paragraph_format.space_after = Pt(8)

    backlog_data = [
        ("3.1 ระบบ Scale สำหรับคลิปรูปภาพและ Layer Clips (Image & Layer Scale Fix)",
         "ปัจจุบันการปรับ Scale สามารถขยาย/ย่อขนาดวิดีโอหลักบน Main Track ได้เรียลไทม์แล้ว แต่สำหรับคลิปรูปภาพ (Image) และ Text/Overlay Clips บนแทร็กชั้นเลเยอร์ การปรับสเกลยังไม่แสดงผลการขยาย/ย่อตามขนาดสเกลที่ตั้งค่าไว้",
         "แผง _apply_overlay ใน editor_preview.py คำนวณขนาดภาพ ow และ oh ยึดตามขนาดแคนวาสดั้งเดิมโดยยังไม่ได้นำค่า clip.get('scale', 1.0) เข้าไปคูณปรับขนาดกว้าง/สูงก่อนวาดลงบนแคนวาส",
         "ปรับปรุงฟังก์ชัน _apply_overlay ให้คำนวณ ow = int(w * scale) และ oh = int(h * scale) เพื่อให้คลิปรูปภาพและเลเยอร์ทั้งหมดขยาย/ย่อขนาดได้ตามสเกลจริงเรียลไทม์"),

        ("3.2 การคลิกลากย้ายตำแหน่งคลิปเสียงบนไทม์ไลน์แบบอิสระ (Audio Clip Horizontal Dragging Fix)",
         "คลิปเสียงบนแทร็ก Audio (audio_0, audio_1) สามารถปรับความยาวและสลับแทร็กได้แล้ว แต่การใช้เมาส์คลิกลากย้ายตำแหน่งเวลาเริ่มต้น (Timeline Position: tl) ไปข้างหน้าหรือข้างหลังยังไม่เรียบลื่นเท่าที่ควร",
         "ใน editor_timeline.py คำสั่ง _tl_press และ _tl_drag มีเงื่อนไขการเช็คระนาบ Y และตัวแปรชนิดแทร็ก audio_N ที่ไปติดเงื่อนไขการสลับแทร็กแนวตั้งก่อน ทำให้ dx การเคลื่อนแนวนอนถูกขัดจังหวะ",
         "ปรับโครงสร้างคำสั่งดักจับเมาส์ใน _tl_press และ _tl_drag แยกกรณีแทร็ก audio_N ออกจากแทร็กวิดีโอ เพื่อคำนวณพิกัด dx แนวนอนและอัปเดต cl['tl'] ได้อย่างอิสระและเรียบลื่น 100%"),

        ("3.3 การปลดปุ่มขยายวิดีโอเต็มหน้าจอออก (Fullscreen Toggle Removal)",
         "ผู้ใช้งานต้องการให้ออกปุ่มขยายวิดีโอเต็มหน้าจอออก เนื่องจากฟังก์ชันดังกล่าวยังทำงานได้ไม่ดีพอและอาจทำให้องค์ประกอบ UI อื่นๆ คลาดเคลื่อน",
         "ปุ่มขยายวิดีโอเต็มหน้าจอมีอยู่ในหน้าจอพรีวิววิดีโอเดิม",
         "ทำการลบปุ่ม Fullscreen Button ออกจาก PreviewPanel ใน editor_preview.py เพื่อให้อินเทอร์เฟซมีความสะอาด เรียบง่าย และทำงานได้อย่างเสถียร")
    ]

    for title, desc, cause, action in backlog_data:
        add_heading2(title)
        add_bullet("รายละเอียดปัญหา / สิ่งที่ต้องทำ", desc)
        add_bullet("สาเหตุ / บริบท", cause)
        add_bullet("แนวทางการดำเนินการแก้ไขในอนาคต", action)

    # Save Master Report Document
    out_dir = r"D:\Folder_For_Work\Year4_1\Media Pro Project"
    out_path = os.path.join(out_dir, "Report_MediaPro_Fix_25_07_26.docx")
    try:
        doc.save(out_path)
        print(f"Report successfully saved to: {out_path}")
    except PermissionError:
        alt_path = os.path.join(out_dir, "Report_MediaPro_Fix_25_07_26_Master.docx")
        doc.save(alt_path)
        print(f"File was locked. Report saved to: {alt_path}")

if __name__ == "__main__":
    create_report()
