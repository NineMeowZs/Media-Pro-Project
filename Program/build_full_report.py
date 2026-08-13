import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Page Margins (1 inch / 2.54 cm all sides)
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Base Font Setup (TH Sarabun New)
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'TH Sarabun New'
    normal_style.font.size = Pt(16)
    normal_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    def set_run_font(run, size=16, bold=False, italic=False, color_rgb=(0x11, 0x18, 0x27)):
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=22, bold=True, color_rgb=(0x0f, 0x17, 0x2a))
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=16, bold=True, color_rgb=(0x33, 0x41, 0x55))
        p.paragraph_format.space_after = Pt(10)
        return p

    def add_heading1(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=18, bold=True, color_rgb=(0x1e, 0x3a, 0x8a))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_heading2(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=16, bold=True, color_rgb=(0x02, 0x84, 0xc7))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        return p

    def add_body(text, bold_prefix="", space_after=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            set_run_font(r_pre, size=16, bold=True, color_rgb=(0x1e, 0x29, 0x3b))
        r_txt = p.add_run(text)
        set_run_font(r_txt, size=16, bold=False, color_rgb=(0x33, 0x41, 0x55))
        return p

    def add_bullet(label, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r_lbl = p.add_run(f"• {label}: ")
        set_run_font(r_lbl, size=15, bold=True, color_rgb=(0x0f, 0x17, 0x2a))
        r_txt = p.add_run(text)
        set_run_font(r_txt, size=15, bold=False, color_rgb=(0x33, 0x41, 0x55))
        return p

    def set_cell_background(cell, fill_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    # ── COVER / HEADER SECTION ──────────────────────────────────────────────
    add_title("แบบฟอร์มบันทึกความก้าวหน้า วิชา MDT371")
    add_subtitle("ชื่อโปรเจกต์: MediaPro พัฒนาโปรแกรมตัดต่อวิดีโอ ที่ใช้ปัญญาประดิษฐ์ในการช่วยทำงาน")

    # Member Table (Table 0)
    t0 = doc.add_table(rows=4, cols=4)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t0, color="CBD5E1", sz="4")
    
    headers_t0 = ["ลำดับที่", "ชื่อ – นามสกุล", "รหัสนักศึกษา", "วิชาเอก"]
    widths_t0 = [Inches(0.8), Inches(2.8), Inches(1.8), Inches(1.1)]

    for c_idx, text in enumerate(headers_t0):
        cell = t0.rows[0].cells[c_idx]
        set_cell_background(cell, "1E293B")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=15, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    members = [
        ("1", "รวิพล ประพันธ์วงศ์", "66120501013", "DMT"),
        ("2", "สุทธิโชติ ม่านทอง", "66120501015", "DMT"),
        ("3", "ธัญกร โชคเกรียงไกร", "66120501056", "DMT"),
    ]

    for r_idx, m in enumerate(members, start=1):
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(m):
            cell = t0.rows[r_idx].cells[c_idx]
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_font(r, size=15, bold=False)

    for row in t0.rows:
        for c_idx, w in enumerate(widths_t0):
            row.cells[c_idx].width = w

    # Metadata Paragraphs
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(12)
    r = p_meta.add_run("อาจารย์ที่ปรึกษา: ดร.วัชระ เรืองสังข์\n")
    set_run_font(r, size=16, bold=True)
    r2 = p_meta.add_run("การเข้าพบอาจารย์ ครั้งที่ 3 วันที่ 27/07/2569 (รายงานสรุปการพัฒนารวบยอดความก้าวหน้าโปรแกรม)")
    set_run_font(r2, size=16, bold=False, color_rgb=(0x02, 0x84, 0xc7))

    # Overview Summary Paragraph
    add_body("ทำการปรับปรุงแก้ไขและอัปเดตประสิทธิภาพการใช้งานของผู้ใช้ (UX) รวมทั้งปรับแต่งหน้าต่างการทำงาน (UI) ของโปรแกรม MediaPro ตามข้อเสนอแนะจากการทดสอบใช้งานจริงของผู้ใช้ โดยรวมรวบยอดหัวข้อการแก้ไขปัญหาที่ผ่านๆ มาทั้งหมด 12 หัวข้อหลัก นำการแก้ไขทางเทคนิคและแนวทางการทำงานแบบไฟนอล (Final Architecture) มาสรุปไว้อย่างเป็นระบบ เพื่อให้โปรแกรมเป็นโปรแกรมตัดต่อวิดีโอระดับมืออาชีพที่ใช้งานง่าย สะดวก รวดเร็ว และลื่นไหลที่สุด", bold_prefix="รายละเอียดความก้าวหน้าและการพัฒนารวบยอด: ")

    doc.add_page_break()

    # ── SECTION 1: REPORT SUMMARY ──────────────────────────────────────────
    add_title("รายงานสรุปการปรับปรุงแก้ไขและพัฒนาโปรแกรม MediaPro (Final Report)")

    add_heading1("1.1 ตารางสรุปภาพรวมการแก้ไขและพัฒนาฟีเจอร์ (Comprehensive Fixes Summary Table)")
    add_body("ตารางด้านล่างนี้แสดงสรุปภาพรวมการดำเนินงานปรับปรุงแก้ไขข้อผิดพลาด (Bug Fixes) และการพัฒนาปรับแต่งฟีเจอร์ (Feature Enhancement) สำหรับโปรแกรม MediaPro โดยรวบรวมรวบยอดปัญหาที่พบ สาเหตุหลัก และแนวทางการแก้ไขทางเทคนิคแบบไฟนอล (Final Solutions) รวมทั้งสิ้น 12 หัวข้อหลัก:")

    # Table 1: Summary Table (12 rows + 1 header)
    t1 = doc.add_table(rows=13, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1, color="94A3B8", sz="4")

    headers_t1 = ["หัวข้อปัญหาที่พบและต้องการแก้ไข", "สาเหตุหลักของปัญหาและความต้องการ", "วิธีการดำเนินการแก้ไขและผลลัพธ์การทำงานแบบไฟนอล"]
    widths_t1 = [Inches(2.2), Inches(2.2), Inches(2.3)]

    for c_idx, text in enumerate(headers_t1):
        cell = t1.rows[0].cells[c_idx]
        set_cell_background(cell, "0F172A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=15, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    summary_data = [
        ("1. การซิงก์ระบบซับไตเติล 2 ทิศทาง และการลบซับ",
         "ลบซับบนไทม์ไลน์แล้วในหมวดซับไม่ลบ หรือลบในหมวดซับแล้วบนไทม์ไลน์ไม่ลบ และกด Ctrl+Z / Ctrl+Y แล้วซับไม่ย้อนกลับพร้อมกัน รวมถึงลากคลุมลบซับไม่ได้",
         "พัฒนาระบบ 2-Way Data Sync สองทิศทางผ่าน _sync_segments_from_tracks() เพิ่มปุ่มถังขยะสีแดง 🗑 Delete All และแก้ไข Multi-Select Delete แบบเรียง Index ย้อนกลับ พร้อมปรับโครงสร้าง Undo/Redo Engine บันทึก Snapshot ก่อนแก้ไข ทำให้ Ctrl+Z/Ctrl+Y คืนค่าซับพร้อมกันทั้งสองหน้าต่าง 100%"),

        ("2. การเร่งความเร็วการประมวลผลด้วย GPU (Hardware Acceleration)",
         "โปรแกรมกินแต่ CPU วิ่ง 80% RAM วิ่ง 70% แต่ GPU วิ่ง 0% ทำให้การประมวลผลและส่งออกวิดีโอช้าและเครื่องกระตุก",
         "พัฒนาระบบ ตรวจจับการ์ดจออัตโนมัติ (NVIDIA NVENC, Intel QSV, AMD AMF) ใน video_exporter.py และ proxy_manager.py ใช้ PyTorch CUDA บน RTX 3060 ถอดเสียง Whisper และเปิด OpenCV OpenCL (cv2.ocl) เร่งความเร็วการประมวลผลเฟรมภาพสด"),

        ("3. การเปลี่ยนไลบรารีและเอนจินพรีวิววิดีโอ (Decord GPU NVDEC Engine)",
         "ไลบรารีเดิม (imageio_ffmpeg / OpenCV CPU Seeking) ใช้คำสั่ง Subprocess Pipe และ Seek รายเฟรม บังคับถอดรหัส Keyframe ใหม่ทำให้พรีวิววิดีโอกระตุกหนัก (10 FPS)",
         "เปลี่ยนไลบรารีถอดรหัสพรีวิวเป็น Decord GPU NVDEC Engine (decord.gpu(0)) ใน video_display_engine.py ถอดรหัสผ่านการ์ดจอ NVIDIA RTX 3060 เข้า VRAM โดยตรง พร้อมอ่านเฟรมแบบ Zero-Seek Sequential Streaming ทำให้พรีวิวเพิ่มขึ้นเป็น 30-60 FPS ลื่นไหลระดับ CapCut"),

        ("4. พิกัดคลิกเลือกคลิปบนไทม์ไลน์และการจัดการ Playhead",
         "กดเลือกคลิปบนไทม์ไลน์ยากมาก พิกัดการกดเหลื่อมลงด้านล่าง (ต้องกดกรอบดำใต้คลิป) และมักกดติด Playhead แทนคลิป",
         "แก้ไขพิกัด Y-Offset ใน Hit-Testing Loops ใน editor_timeline.py โดยข้าม __empty_layer__ ทำให้พิกัดกดตรงกับบล็อกคลิป 100% จัดลำดับ Hitbox ให้คลิปมีความสำคัญก่อน Playhead จำกัดพื้นที่จับ Playhead เฉพาะหัวสามเหลี่ยม และเพิ่ม Edge Snapping หยุดชะงักตามขอบคลิป"),

        ("5. ระบบปรับแต่ง Text & Subtitle Properties Panel",
         "ซับไตเติลปรับฟอนต์/สไตล์ไม่เปลี่ยนตาม Slider ขนาดปรับยาก และกดเลือกซับในรายการแล้วแสดงการตั้งค่าผิดอัน",
         "ออกแบบ Properties Panel รวมศูนย์ใน editor_properties.py เปลี่ยน Slider เป็น CTkEntry กรอกขนาดพิกเซล (px) โดยตรง รองรับฟอนต์ไทย สไตล์ (Outline, Shadow, Box) แอนิเมชัน และเพิ่มปุ่ม 'Apply to All Subtitles' ปรับใช้กับซับทั้งหมดในคลิกเดียว พร้อมแก้ Index ซับที่เลือกให้ตรงกับจริง"),

        ("6. ระบบประมวลผลคลื่นเสียงและการแยกเสียงวิดีโอ (Audio & Waveform)",
         "ไม่เห็นคลื่นเสียงบนไทม์ไลน์ ไม่มีเสียงในคลิปพรีวิว และต้องการแยกเสียงออกจากคลิปวิดีโอเฉพาะส่วนที่ตัด",
         "สกัดค่าแอมพลิจูดความถี่เสียงวาด Waveform บนแทร็กวิดีโอ/เสียง เพิ่มฟังก์ชัน Detach Audio ในคลิกขวาแยกเสียงเฉพาะท่อนที่ Split ออกเป็นคลิปเสียงบน Audio 1 พร้อม Mute วิดีโอหลัก และปรับปรุง _setup_audio() ผสมเสียงสดทุกคลิปบนไทม์ไลน์เล่นผ่าน Pygame Mixer เต็มคุณภาพ"),

        ("7. ระบบสร้างซับไตเติลตามสปีดวิดีโอ & VAD Silence Deadair Filtering",
         "ซับไตเติลไปไวและมั่วช่วงเสียงเงียบ และเมื่อวิดีโอถูกเร่งสปีด/สโลว์ ซับไตเติลยังอิงตามสปีดวิดีโอดั้งเดิม",
         "พัฒนาระบบ Voice Activity Detection (VAD) แบบ RMS Energy ตัดขอบเขตช่วงเดดแอร์ (>0.35s) ออก พร้อมคำนวณกระจายเวลาตามจำนวนอักขระ และเพิ่ม _build_atempo_filter() ปรับสเกลเวลาเสียงตามสปีดคลิป (0.1x-10.0x) ทำให้ซับตรงกับเสียงพูดที่เร่ง/สโลว์ และทำซับครอบคลุมทุกวิดีโอทั้งโปรเจกต์"),

        ("8. การปรับโครงสร้างเลเยอร์ไทม์ไลน์แบบไดนามิก (Clean Dynamic Layout)",
         "เลเยอร์ไทม์ไลน์รกรุงรัง ซูมเข้าออกอัตโนมัติขณะเลื่อนคลิปจนมึนหัว และการ Scroll เมาส์เลื่อนแนวนอนใช้งานไม่สะดวก",
         "ปรับหน้าตาไทม์ไลน์ใหม่ เริ่มต้นแสดงเฉพาะ Main Video, Subtitle และ Audio 1 ใช้ไอคอนสัญลักษณ์ขนาดใหญ่ (▶, T, ♫, ■) แทนชื่อเลเยอร์ รองรับการเพิ่มเลเยอร์ Audio/Overlay ไดนามิกเมื่อลากคลิปมาวาง ปิด Auto-Zoom ปรับ Mouse Scroll เป็นเลื่อนแนวตั้ง และใช้ Ctrl+Wheel สำหรับซูม"),

        ("9. การปรับกระบวนการสร้างโปรเจกต์เปล่า (Blank Project Initialization)",
         "การสร้างโปรเจกต์ใหม่บังคับให้เลือกไฟล์วิดีโอตั้งต้นทันที ทำให้ทำซับไม่อิงตามไฟล์วิดีโอจริงที่นำมาตัดต่อ",
         "ปรับปุ่ม ＋ Create Project ในหน้าแรก (app.py) ให้เปิดหน้าจอโปรเจกต์ว่างเปล่า (EditorPage(initial_video=None)) ทันที โดยผู้ใช้สามารถนำเข้าสื่อเข้า Asset Library แล้วลากวางลงไทม์ไลน์ได้อย่างอิสระ"),

        ("10. ระบบส่งออกวิดีโอไร้ Noise คมชัดสูง และแสดง Progress %",
         "ตอนส่งออกวิดีโอไม่มีเปอร์เซ็นต์ความคืบหน้า ชื่อไฟล์ไม่ตรงชื่อโปรเจกต์ และไฟล์วิดีโอส่งออกมี Noise และรอยบีบอัด",
         "อ่านค่า stderr จาก FFmpeg Popen แสดงเปอร์เซ็นต์ความคืบหน้า 0%-100% บน Status bar ตั้งชื่อไฟล์ผลลัพธ์เป็น {ชื่อโปรเจกต์}.mp4 อัตโนมัติ และอัปเกรดพารามิเตอร์ NVENC HQ 2-Pass VBR (-preset p6 -rc vbr -cq 17 -b:v 16M -spatial-aq 1) ขจัด Noise คมชัดสูงสุด"),

        ("11. คีย์ลัดตัดคลิป Ctrl+B และการปรับความยาวคลิปสองฝั่ง",
         "ไม่มีคีย์ลัด Ctrl+B สำหรับตัดคลิป และข้อความ/ซับไตเติลดึงขยายความยาวได้เฉพาะจากด้านหลัง เพิ่มจากด้านหน้าไม่ได้",
         "เพิ่มคีย์ลัด Ctrl+B ตัดแบ่งคลิป ณ ตำแหน่ง Playhead ทุกองค์ประกอบ และปรับปรุงฟังก์ชัน Trim (trim_l / trim_r) ให้คลิปประเภท Unlimited (Text, Subtitles, Image, Audio) สามารถดึงขยายหรือดันหดขอบเขตได้ทั้งด้านหน้า (ซ้าย) และด้านหลัง (ขวา) อิสระ"),

        ("12. การปรับปรุงความเสถียรไฟล์ SRT และการประมวลผลแบบ Non-Blocking",
         "นำเข้า/ทำไฟล์ SRT ไม่ได้ เกิด Unicode Error และคอมพิวเตอร์ค้างตอบสนองขณะประมวลผลทำซับไตเติล",
         "ปรับปรุงการอ่านเขียนไฟล์ SRT ให้รองรับ UTF-8 / UTF-8 with BOM อัตโนมัติใน transcriber.py และย้ายกระบวนการถอดเสียง AI และการสร้างภาพ Proxy ออกไปทำใน Background Thread ทั้งหมด ทำให้หน้าจอหลักไม่ค้างตอบสนอง (Zero UI Freezes)"),
    ]

    for r_idx, row_data in enumerate(summary_data, start=1):
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = t1.rows[r_idx].cells[c_idx]
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_font(r, size=14, bold=(c_idx == 0))

    for row in t1.rows:
        for c_idx, w in enumerate(widths_t1):
            row.cells[c_idx].width = w

    doc.add_page_break()

    # ── SECTION 1.2: DETAILED TECHNICAL SPECIFICATIONS ─────────────────────
    add_heading1("1.2 รายละเอียดการปรับปรุงแก้ไขและพัฒนาเจาะจง (Detailed Technical Specifications)")
    add_body("ในส่วนนี้เป็นการอธิบายรายละเอียดการดำเนินงานปรับปรุงแก้ไขและพัฒนาฟีเจอร์อย่างเจาะจงทั้ง 12 หัวข้อ โดยครอบคลุมถึงปัญหาที่พบเจอ สาเหตุหลักทางเทคนิค วิธีการแก้ไขเชิงซอฟต์แวร์สถาปัตยกรรม และผลลัพธ์การทำงานแบบไฟนอล ดังนี้:")

    # Detailed Item 1
    add_heading2("1) การซิงก์ระบบซับไตเติล 2 ทิศทาง และการลบซับ (Unified Subtitle Synchronization & Multi-Delete)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "เดิมการลบซับไตเติลบนไทม์ไลน์ไม่ได้เชื่อมโยงกับรายการในหมวดซับไตเติล (Transcript Panel) และเมื่อกดลบในหมวดซับ ไทม์ไลน์กลับไม่ถูกลบตาม ส่งผลให้ข้อมูลซับไตเติลคลาดเคลื่อน นอกจากนี้การกด Ctrl+Z ย้อนกลับการทำงานยังคืนค่าซับไม่พร้อมกัน และการลากคลุมลบซับไตเติลแบบ Multi-Select ไม่สามารถลบซับออกได้เนื่องจากดัชนีรายการเกิดการคลาดเคลื่อน (Index Corruption)")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "พัฒนาระบบซิงก์ข้อมูลแบบ 2-Way Sync โดยสร้างฟังก์ชัน _sync_segments_from_tracks() ใน editor_page.py เพื่อเชื่อมโยง self.tracks['subtitle'] กับ self.segments เป็นออบเจกต์เดียวกัน ปรับปรุงฟังก์ชัน _del_sel() ให้เก็บดัชนีคลิปที่เลือกทั้งหมดจาก _multi_sel แล้วทำการลบเรียงตามดัชนีย้อนกลับ (Descending Order) พร้อมเพิ่มปุ่มถังขยะสีแดง 🗑 Delete All ใน transcript_panel.py สำหรับลบซับทั้งหมดในคลิกเดียว และปรับโครงสร้าง _push_undo(), _undo_do(), _redo_do() ให้บันทึก Snapshot ก่อนการแก้ไข ทำให้ Ctrl+Z และ Ctrl+Y คืนค่าและทำซ้ำซับไตเติลได้พร้อมกันทั้งสองหน้าต่าง")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "การจัดการซับไตเติลมีความแม่นยำและเสถียร 100% การลบซับไตเติลไม่ว่าจะลบบนไทม์ไลน์ ลบในหน้าต่าง Transcript ลบแบบลากคลุมหลายรายการ หรือลบทั้งหมดด้วยปุ่มถังขยะสีแดง ข้อมูลจะถูกซิงก์ลบออกตรงกันทันที และการกด Ctrl+Z / Ctrl+Y สามารถย้อนกลับและทำซ้ำการลบซับได้สมบูรณ์แบบโดยไม่มีข้อมูลตกค้าง")

    # Detailed Item 2
    add_heading2("2) การเร่งความเร็วการประมวลผลด้วย GPU (Hardware Acceleration via NVENC, PyTorch CUDA & OpenCL)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "การประมวลผลภาพพรีวิว การสร้างไฟล์ Proxy การถอดเสียงด้วย AI และการส่งออกวิดีโอ พึ่งพาการทำงานของ CPU เป็นหลัก (CPU วิ่ง 80-90% RAM วิ่ง 70% แต่ GPU วิ่ง 0%) ส่งผลให้เครื่องเกิดความร้อนสูง ทำงานได้ช้า และหน้าจอพรีวิวเกิดอาการกระตุกอย่างหนัก")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "พัฒนาระบบ Hardware Acceleration ครอบคลุมทั้งโปรแกรม โดยสร้างฟังก์ชัน _detect_gpu_encoder() ใน video_exporter.py และ proxy_manager.py ตรวจจับฮาร์ดแวร์การ์ดจออัตโนมัติ (NVIDIA NVENC h264_nvenc, Intel QSV h264_qsv, AMD AMF h264_amf) สำหรับการสร้าง Proxy และส่งออกวิดีโอ ในส่วนโมเดล Whisper ถอดเสียงภาษาไทย กำหนดให้โหลดน้ำหนักเข้าสู่หน่วยความจำ VRAM ของ NVIDIA GeForce RTX 3060 Laptop GPU ผ่าน PyTorch CUDA และเปิดใช้งาน OpenCV OpenCL (cv2.ocl.setUseOpenCL(True)) สำหรับประมวลผลเฟรมภาพสด")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "การ์ดจอ GPU ถูกดึงมาช่วยประมวลผลอย่างเต็มประสิทธิภาพ GPU Utilization เพิ่มขึ้นเป็น 40-70% ขณะเรนเดอร์และส่งออกวิดีโอ ช่วยลดภาระการทำงานของ CPU ลงอย่างเห็นได้ชัด ความเร็วในการสร้าง Proxy และ Export วิดีโอเพิ่มขึ้นกว่า 3-5 เท่า")

    # Detailed Item 3
    add_heading2("3) การเปลี่ยนไลบรารีและเอนจินถอดรหัสพรีวิววิดีโอ (Decord GPU NVDEC Hardware Engine)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "การแสดงผลพรีวิววิดีโอบนไทม์ไลน์เดิมมีความกระตุกอย่างมาก เฟรมเรตตกเหลือเพียง 10 FPS สาเหตุหลักเกิดจากไลบรารีและกระบวนการพรีวิวเดิมใช้งาน imageio_ffmpeg สปอว์นกระบวนการ Subprocess FFmpeg Pipe รายเฟรม หรือใช้ cv2.VideoCapture.set(CAP_PROP_POS_MSEC) ถอดรหัสผ่าน CPU ซึ่งบังคับให้ระบบต้องค้นหา Keyframe และถอดรหัสเฟรมกลางใหม่ตั้งแต่ต้นทุกครั้ง รวมทั้งเกิด Latency จากการโอนย้าย Tensor ระหว่าง CPU และ GPU ข้ามไปมา")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "เปลี่ยนไลบรารีถอดรหัสพรีวิววิดีโอหลักมาเป็น Decord (High-Performance GPU Video Loader) โดยสร้างสถาปัตยกรรม SmartVideoReader ใน video_display_engine.py เรียกใช้ decord.gpu(0) ดึงฮาร์ดแวร์ NVIDIA NVDEC บนการ์ดจอ RTX 3060 ถอดรหัสไฟล์วิดีโอตรงเข้าสู่หน่วยความจำ VRAM พร้อมทั้งเพิ่มฟังก์ชัน read_next_frame() อ่านเฟรมแบบ Zero-Seek Sequential Streaming ขณะเล่นวิดีโอต่อเนื่อง และขจัด Bottleneck การโอนย้าย Tensor ข้าม bus ด้วย OpenCV OpenCL C++ Accelerated Resizing")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "การเปลี่ยนมาใช้ไลบรารี Decord GPU NVDEC ร่วมกับเอนจิน Zero-Seek Streaming ทำให้อัตราเฟรมพรีวิวตัดต่อวิดีโอเพิ่มขึ้นจากเดิม 10 FPS เป็น 30-60 FPS สมูธ นุ่มนวล ไม่มีอาการกระตุกหรือเฟรมกระตุกสะดุด มอบประสบการณ์ตัดต่อวิดีโอระดับเดียวกับโปรแกรมชั้นนำอย่าง CapCut")

    # Detailed Item 4
    add_heading2("4) การปรับปรุงพิกัดคลิกเลือกคลิปบนไทม์ไลน์และการจัดการ Playhead (Hit-Test Alignment & Edge Snapping)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "ผู้ใช้ประสบปัญหาอย่างหนักในการคลิกเลือกบล็อกคลิปบนไทม์ไลน์ โดยพิกัดการคลิกเกิดการเหลื่อมลงด้านล่าง (ผู้ใช้ต้องการกดเลือกคลิปในกรอบสีน้ำเงิน แต่ต้องไปกดตรงกรอบดำด้านล่างจึงจะติด) และเมื่อกดคลิกมักจะไปโดน Playhead แทนบล็อกคลิป นอกจากนี้ Playhead ยังเลื่อนผ่านขอบคลิปได้ยาก ขาดจุดหยุดชะงัก")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ค้นพบสาเหตุหลักว่าลูปตรวจจับตำแหน่งเมาส์ (Hit-Testing Loops) ใน editor_timeline.py คำนวณบวกความสูงของ __empty_layer__ เข้าไปในค่า Y ในขณะที่ลูปวาดภาพข้ามแถบว่าง ทำให้พิกัดการกดเหลื่อมลงล่าง 1 แถวเต็มๆ จึงทำการใส่เงื่อนไข if key == '__empty_layer__': continue ให้ครอบคลุมทุกลูปเมาส์ (_tl_press, _tl_hover, _tl_rclick, _tl_drag, _tl_release) ปรับลำดับ Hitbox ให้บล็อกคลิปมีความสำคัญอันดับแรก (Priority 1) จำกัดพื้นที่จับ Playhead เฉพาะบริเวณหัวสามเหลี่ยม (Ruler Cap) และเพิ่มระบบ Snap Playhead หยุดชะงักตามขอบคลิป (Threshold 8px)")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "การคลิกเลือกบล็อกคลิปวิดีโอ ซับไตเติล เสียง และรูปภาพบนไทม์ไลน์ มีความแม่นยำตรงกับตำแหน่งที่มองเห็นบนหน้าจอ 100% (คลิกในกรอบสีน้ำเงินติดทันที) ไม่เกิดการกดวืดหรือติด Playhead โดยไม่ตั้งใจ และ Playhead มีระบบ Snap ขอบคลิปช่วยให้ควบคุมตำแหน่งตัดต่อได้ง่ายและแม่นยำ")

    # Detailed Item 5
    add_heading2("5) ระบบปรับแต่ง Text & Subtitle Properties Panel (Numeric Pixel Entries & Apply to All)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "พาเนลปรับแต่ง Properties ของซับไตเติลและข้อความไม่เปลี่ยนฟอนต์และสไตล์ตามที่เลือก การปรับขนาดฟอนต์ด้วย Slider ควบคุมตัวเลขพิกเซลที่แม่นยำได้ยาก และเมื่อกดเลือกซับไตเติลในรายการ Transcript หน้าต่าง Properties กลับแสดงผลและแก้ไขซับคลาดเคลื่อนผิดอัน")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ออกแบบ Properties Panel รวมศูนย์ใน editor_properties.py ให้รองรับทั้ง Text Overlays และ Subtitle Clips ยกเลิกการใช้ Slider เปลี่ยนเป็น CTkEntry กรอกขนาดฟอนต์เป็นตัวเลขพิกเซล (px) โดยตรง เพิ่ม Dropdown เลือกฟอนต์ภาษาไทยมาตรฐาน (Tahoma, TH Sarabun, Arial ฯลฯ) เลือกสไตล์ (Outline, Shadow, Box) แอนิเมชัน และตำแหน่ง เพิ่มปุ่ม 'Apply to All Subtitles' ปรับใช้การตั้งค่ากับซับไตเติลทั้งหมดในโปรเจกต์พร้อมกัน และแก้ไขดัชนี sel_idx ให้เชื่อมโยงกับดัชนี Segment ที่เลือกใน Transcript Panel อย่างถูกต้อง")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "ผู้ใช้สามารถปรับแต่งฟอนต์ ขนาดสเกลพิกเซล สี สไตล์ และแอนิเมชันของข้อความและซับไตเติลได้อย่างแม่นยำ ฟอนต์และสไตล์เปลี่ยนตามที่เลือกทันที สามารถกดปุ่มปรับใช้กับซับทั้งหมดได้ในคลิกเดียว และการกดเลือกซับในหน้าต่าง Transcript แสดงผลตรงกับออบเจกต์ที่ต้องการแก้ไข 100%")

    # Detailed Item 6
    add_heading2("6) ระบบประมวลผลคลื่นเสียงและการแยกเสียงวิดีโอ (Audio Waveform, Detach Audio & Preview Audio Synthesis)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "แทร็กวิดีโอและเสียงบนไทม์ไลน์เป็นแถบสีเรียบ มองไม่เห็นช่วงที่มีเสียงพูดหรือเสียงเงียบ คลิปเสียงที่นำเข้ามาไม่มีเสียงเล่น และไม่มีฟังก์ชันแยกเสียงออกจากคลิปวิดีโอ รวมทั้งพรีวิววิดีโอไม่มีเสียงเล่นขณะตัดต่อ")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "พัฒนาระบบสกัดค่าแอมพลิจูดความถี่เสียงด้วย FFmpeg/NumPy ใน Background Thread วาดกราฟคลื่นเสียง Waveform ลงบนบล็อกคลิปวิดีโอและคลิปเสียงบนไทม์ไลน์ เพิ่มฟังก์ชัน '🔊 Detach Audio' ใน Context Menu (คลิกขวา) ทำการแยกเสียงเฉพาะท่อนที่ถูก Split ออกมาสร้างเป็นคลิปเสียงใหม่บนแทร็ก Audio 1 พร้อมตั้งค่า Mute ปิดเสียงแทร็กวิดีโอเดิมเพื่อป้องกันเสียงซ้อน และปรับปรุงระบบผสมเสียงพรีวิว _setup_audio() ใน editor_page.py รวมสัญญาณเสียงสดจากทุกคลิปบนไทม์ไลน์ส่งให้ Pygame Mixer เล่นเสียงพรีวิวขณะตัดต่อ")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "คลิปวิดีโอและคลิปเสียงบนไทม์ไลน์แสดงกราฟคลื่นเสียง Waveform ชัดเจน ช่วยให้ผู้ใช้ตัดต่อเข้าจังหวะเสียงได้ง่าย สามารถแยกเสียงออกจากวิดีโอท่อนที่ Split ได้สมบูรณ์ และพรีวิววิดีโอเล่นเสียงได้เต็มคุณภาพ ตรงตามตำแหน่งเวลาบนไทม์ไลน์ตลอดการตัดต่อ")

    doc.add_page_break()

    # Detailed Item 7
    add_heading2("7) ระบบสร้างซับไตเติลตามสปีดวิดีโอ & VAD Silence Deadair Filtering (Speed-Adapted Auto Subtitles)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "การสร้างซับไตเติลเดิมมีปัญหาซับขึ้นไวและมั่วช่วงเสียงเงียบเนื่องจากไม่มีการตรวจจับเสียงพูดจริง และเมื่อวิดีโอถูกปรับความเร็วสปีด (เร่งความเร็ว 2x หรือสโลว์โมชัน 0.5x) ซับไตเติลที่สร้างออกมายังคงอิงตามสปีดวิดีโอดั้งเดิม 1.0x ทำให้ซับหลุดจังหวะไม่ตรงกับเสียงพูด รวมทั้งถอดเสียงได้เฉพาะคลิปแรก ไม่ครอบคลุมทุกวิดีโอบนไทม์ไลน์")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "พัฒนาระบบ Voice Activity Detection (VAD) แบบ RMS Audio Energy ใน transcriber.py ตรวจวัดระดับพลังงานเสียงและตัดขอบเขตช่วงเดดแอร์/เสียงเงียบ (>0.35s) ออก พร้อมคำนวณการกระจายความยาวตามจำนวนอักขระ (Character-Weighted Duration Allocation) ควบคู่กับ Whisper Timestamps สร้างฟังก์ชัน _build_atempo_filter() ใน editor_page.py ปรับแต่งสเกลเวลาเสียงตามค่า Speed (0.1x-10.0x) ของทุกคลิปวิดีโอบน Main Track รวมออกมาเป็นไฟล์ combined_main_audio.wav ส่งให้ Whisper ถอดเสียง")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "ซับไตเติลที่สร้างขึ้นมีความแม่นยำสูง ไม่ขึ้นมั่วในช่วงเสียงเงียบ เมื่อคลิปวิดีโอถูกเร่งสปีดหรือสโลว์โมชัน ซับไตเติลจะปรับสเกลเวลาตามความเร็ววิดีโอให้อัตโนมัติ จังหวะซับตรงกับเสียงพูดบนวิดีโอ 100% และรองรับการทำซับครอบคลุมทุกคลิปวิดีโอทั่วทั้งโปรเจกต์ (Full Project Transcription)")

    # Detailed Item 8
    add_heading2("8) การปรับโครงสร้างเลเยอร์ไทม์ไลน์แบบไดนามิก (Clean Dynamic Layout & Vertical Scroll)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "เลเยอร์ไทม์ไลน์เดิมมีแทร็ก Music และ Audio 2 แสดงค้างไว้รกรุงรัง การขยับคลิปออกนอกพื้นที่สายตาทำให้ไทม์ไลน์ซูมเข้าออกอัตโนมัติจนเกิดอาการมึนหัว และการ Scroll เมาส์กลายเป็นการเลื่อนไทม์ไลน์แนวนอนซึ่งไม่คุ้นเคยต่อผู้ใช้งาน")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ปรับดีไซน์ไทม์ไลน์ใหม่ให้เรียบหรู โดยเริ่มต้นแสดงเฉพาะ Main Video, Subtitle และ Audio 1 เปลี่ยนป้ายเลเยอร์เป็นไอคอนสัญลักษณ์ขนาดใหญ่ทันสมัย (▶ Video, T Subtitle, ♫ Audio, ■ Overlay) รองรับการสร้างเลเยอร์เสียง (Audio 2, Audio 3...) และเลเยอร์วิดีโอซ้อน (Layer 1, Layer 2...) อัตโนมัติเมื่อลากคลิปมาวาง ยึดระดับ Scale ไทม์ไลน์ให้คงที่ขณะลากคลิป (ปิด Auto-Zoom) และปรับ Mouse Scroll เป็นการเลื่อนแนวดิ่ง (Vertical Scroll) โดยใช้ Ctrl+Wheel สำหรับซูมไทม์ไลน์")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "หน้าต่างไทม์ไลน์มีความสะอาด สวยงาม ดูเป็นมืออาชีพ การลากวางคลิปไม่เกิดการซูมวูบวาบ มึนหัว และการ Scroll เมาส์เลื่อนดูเลเยอร์แนวดิ่งใช้งานได้สะดวก ราบรื่น เป็นธรรมชาติ")

    # Detailed Item 9
    add_heading2("9) การปรับกระบวนการสร้างโปรเจกต์เปล่า (Blank Project Initialization Flow)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "เดิมการกดสร้างโปรเจกต์ใหม่ (Create Project) บังคับให้ผู้ใช้ต้องเลือกไฟล์วิดีโอตั้งต้นทันที ส่งผลให้เวลาทำซับไตเติลหรือตัดต่อเกิดการยึดติดกับไฟล์วิดีโอดั้งเดิมไฟล์แรก ไม่สามารถสร้างโปรเจกต์ว่างเพื่อนำเข้าไฟล์วิดีโอหลายๆ ไฟล์ได้อย่างเป็นอิสระ")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ปรับเปลี่ยนคำสั่งใน HomePage (app.py) บริเวณการกดปุ่ม ＋ Create Project ให้ทำการเปิดหน้าต่างตัดต่อโดยส่งค่า initial_video=None เข้าสู่ออบเจกต์ EditorPage โดยตรง ทำให้เกิดการสร้างโปรเจกต์ว่างเปล่า (Blank Project) ที่มีไทม์ไลน์และ Asset Library ว่างพร้อมใช้งาน")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "ผู้ใช้สามารถเริ่มต้นสร้างโปรเจกต์เปล่าได้อย่างอิสระ แล้วค่อยนำเข้าวิดีโอ รูปภาพ หรือเสียงหลายๆ ไฟล์เข้าสู่ Asset Library และลากวางลงบนไทม์ไลน์เพื่อตัดต่อและทำซับไตเติลอิงจากวิดีโอที่ใช้งานจริงบนไทม์ไลน์ได้อย่างสมบูรณ์")

    # Detailed Item 10
    add_heading2("10) ระบบส่งออกวิดีโอไร้ Noise คมชัดสูง และแสดงเปอร์เซ็นต์ความคืบหน้า (Noise-Free HQ Export & Live Progress %)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "ขณะเรนเดอร์ส่งออกวิดีโอไม่มีการแสดงเปอร์เซ็นต์ความคืบหน้า ชื่อไฟล์ส่งออกค้างเป็นค่าว่างหรือ Output_Video.mp4 และไฟล์วิดีโอที่ได้มีปัญหา Noise/รอยบีบอัดภาพเป็นบล็อกเนื่องจากการกำหนด Bitrate Cap ของ NVENC ไว้ต่ำเกินไป (6 Mbps)")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "พัฒนาโครงสร้างอ่านค่า stderr จากกระบวนการ FFmpeg Popen แบบ Real-time แปลงค่าสแตมป์เวลาเป็นเปอร์เซ็นต์ความคืบหน้า 0%-100% แสดงผลบน Status bar ส่งค่า proj_name เข้าหน้าต่าง ExportDialog ตั้งชื่อไฟล์เริ่มต้นเป็น {ชื่อโปรเจกต์}.mp4 อัตโนมัติ และอัปเกรดพารามิเตอร์ส่งออกด้วย NVENC HQ 2-Pass VBR Mode (-preset p6 -rc vbr -cq 17 -b:v 16M -maxrate 30M -bufsize 30M -spatial-aq 1 -temporal-aq 1 -pix_fmt yuv420p)")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "ผู้ใช้รับรู้ความคืบหน้าการเรนเดอร์วิดีโอแบบ Real-time 0%-100% ชื่อไฟล์ตั้งให้อัตโนมัติตามโปรเจกต์ และไฟล์วิดีโอผลลัพธ์ MP4 ที่ส่งออกมีความคมชัดสูง ภาพเนียนใส ไร้ปัญหา Noise หรือรอยภาพแตกบีบอัด")

    # Detailed Item 11
    add_heading2("11) คีย์ลัดตัดคลิป Ctrl+B และการปรับความยาวคลิปสองฝั่ง (Ctrl+B Split Shortcut & Dual Trimming)")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "ขาดคีย์ลัดมาตรฐาน Ctrl+B ในการตัดแบ่งคลิป และคลิปประเภทข้อความ ซับไตเติล รูปภาพ และเสียง ไม่สามารถดึงขยายความยาวจากด้านหน้า (ด้านซ้าย) ได้ ดึงขยายได้เฉพาะจากด้านหลังเท่านั้น")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ผูกคีย์ลัด Ctrl+B เข้ากับฟังก์ชัน _split() ใน editor_page.py สำหรับตัดแบ่งคลิป ณ ตำแหน่ง Playhead และปรับปรุงโหมดการ Trim (trim_l และ trim_r) ใน editor_timeline.py สำหรับคลิปประเภท Unlimited (Text, Subtitle, Image, Audio) ให้คำนวณปรับค่า tl, start และ end สอดคล้องกัน ทำให้ดึงขยายหรือดันหดขอบเขตคลิปได้ทั้งฝั่งซ้ายและฝั่งขวา")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "ผู้ใช้สามารถกด Ctrl+B ตัดแบ่งคลิปได้รวดเร็วทันใจ และสามารถดึงขยายปรับความยาวของซับไตเติล ข้อความ รูปภาพ และเสียงบนไทม์ไลน์ได้ทั้งจากด้านหน้าและด้านหลังได้อย่างเป็นอิสระ")

    # Detailed Item 12
    add_heading2("12) การปรับปรุงความเสถียรไฟล์ SRT และการประมวลผลแบบ Non-Blocking Thread")
    add_bullet("ปัญหาที่พบและสาเหตุหลัก", "เกิดข้อผิดพลาดในการทำไฟล์ซับไตเติลและนำเข้าไฟล์ SRT (UnicodeDecodeError) และการประมวลผลถอดเสียง AI หรือการสร้างไฟล์ Proxy ทำให้โปรแกรมเกิดอาการค้าง ไม่ตอบสนอง (UI Freezes)")
    add_bullet("วิธีการแก้ไขและซอฟต์แวร์สถาปัตยกรรม", "ปรับปรุงการอ่านและบันทึกไฟล์ SRT ใน transcriber.py ให้รองรับการเข้ารหัส UTF-8 และ UTF-8 with BOM อัตโนมัติ ย้ายกระบวนการประมวลผลถอดเสียง AI โมเดล Whisper การสกัด Waveform และการสร้างไฟล์ Proxy ออกไปทำงานบน Background Worker Thread ทั้งหมด")
    add_bullet("ผลลัพธ์การทำงานแบบไฟนอล", "การนำเข้าและส่งออกไฟล์ SRT ทำงานได้อย่างแม่นยำ ไม่เกิดข้อผิดพลาดภาษาไทย และโปรแกรม MediaPro ทำงานได้อย่างลื่นไหล เสถียร หน้าจอหลักไม่เกิดอาการค้างตอบสนองแม้ในระหว่างการประมวลผลหนัก")

    doc.add_page_break()

    # ── SECTION 2: IRB STATUS ──────────────────────────────────────────────
    add_heading1("2. ความคืบหน้าของเอกสารที่ต้องส่ง IRB (Human Research Ethics Protocol)")
    add_body("ตารางด้านล่างนี้แสดงสถานะความคืบหน้าในการจัดเตรียมเอกสารจริยธรรมการวิจัยในมนุษย์ (IRB) สำหรับโครงการวิจัยการพัฒนาโปรแกรมตัดต่อวิดีโอ MediaPro:")

    # Table 2: IRB Status Table
    t2 = doc.add_table(rows=5, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2, color="94A3B8", sz="4")

    headers_t2 = ["ลำดับ", "รายการเอกสาร", "มี", "ไม่มี (โปรดชี้แจง)"]
    widths_t2 = [Inches(0.8), Inches(3.6), Inches(0.8), Inches(2.3)]

    for c_idx, text in enumerate(headers_t2):
        cell = t2.rows[0].cells[c_idx]
        set_cell_background(cell, "1E293B")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=15, bold=True, color_rgb=(0xFF, 0xFF, 0xFF))

    irb_data = [
        ("1", "แบบตรวจสอบโครงการวิจัยที่เข้าข่ายการขอประเมินจริยธรรมการวิจัยในมนุษย์ (IRB Checklist)", "✓", "-"),
        ("2", "บันทึกข้อความ และแบบเสนอเพื่อขอประเมินจริยธรรมการวิจัยในมนุษย์ แบบยกเว้น (IRB Form-01)", "-", "อยู่ระหว่างจัดทำแบบสอบถามหลังใช้งานโปรแกรม"),
        ("3", "โครงการวิจัย/กิจกรรมฉบับสมบูรณ์ (Full Proposal Document)", "✓", "-"),
        ("4", "ประกาศนียบัตรการอบรมจริยธรรมการวิจัยในมนุษย์ของ ผู้วิจัย และอาจารย์ที่ปรึกษา", "-", "อยู่ระหว่างการรวบรวมใบประกาศนียบัตร"),
    ]

    for r_idx, row_data in enumerate(irb_data, start=1):
        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = t2.rows[r_idx].cells[c_idx]
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_run_font(r, size=15, bold=False)

    for row in t2.rows:
        for c_idx, w in enumerate(widths_t2):
            row.cells[c_idx].width = w

    add_body("สรุปสถานะเอกสาร IRB: ในปัจจุบันเอกสารคำขอประเมินอยู่ระหว่างการรวบรวมผลแบบสอบถามประเมินความพึงพอใจหลังการทดลองใช้งานโปรแกรม และเตรียมจัดส่งเอกสารให้คณะกรรมการจริยธรรมการวิจัยต่อไป", bold_prefix="\nหมายเหตุสรุป: ")

    # ── SECTION 3: AI EVALUATION ───────────────────────────────────────────
    add_heading1("3. การทดสอบประเมินผล AI Model (CER & Inference Speed Evaluation)")
    add_body("ผลการทดสอบการถอดเสียงภาษาไทยด้วยโมเดล Custom Whisper PyTorch (best.pt / Small Model) เปรียบเทียบกับโมเดลเวอร์ชันเดิม:")
    add_bullet("Character Error Rate (CER)", "อัตราความผิดพลาดระดับอักขระ (CER) ลดลงอย่างมีนัยสำคัญ ปัญหาการหลุดหายของคำบางคำและคำเพี้ยนลดลงอย่างชัดเจน")
    add_bullet("Inference Time Efficiency", "เวลาที่ใช้ในการประมวลผลถอดเสียง (Inference Time) ลดลงอย่างมาก เมื่อรันบน NVIDIA GeForce RTX 3060 Laptop GPU ผ่าน PyTorch CUDA")
    add_bullet("การเปรียบเทียบโมเดล Small เดิม vs โมเดล Small ใหม่", "โมเดล Small เวอร์ชันใหม่ให้ผลลัพธ์อ่านง่าย ได้ความหมายตรงตามบริบท แม้จะมีคำผิดเล็กน้อยในบางคำประธานเฉพาะ แต่องค์รวมการสร้างซับไตเติลนำไปใช้งานตัดต่อจริงได้อย่างมีประสิทธิภาพสูง")

    # ── SECTION 4: FUTURE ROADMAP ──────────────────────────────────────────
    add_heading1("4. รายการปัญหาที่พบเจอและแนวทางการพัฒนาแก้ไขในอนาคต (Future Roadmap)")
    add_body("จากการทดสอบและรวบรวมข้อคิดเห็นเพิ่มเติม คณะผู้จัดทำได้วางแผนแนวทางการพัฒนาและปรับปรุงโปรแกรม MediaPro ต่อเนื่องในอนาคต 2 ประเด็นหลัก ดังนี้:")

    add_bullet("1. ระบบพรีวิว Asset รูปภาพและเพลงสดใน Properties Panel และการ Render Overlays",
               "ปัญหาที่พบ: เมื่อนำเข้าไฟล์รูปภาพและไฟล์เพลงอิสระ พรีวิวสดใน Properties Panel ยังไม่แสดงภาพตัวอย่าง และเมื่อกดส่งออกวิดีโอ (Export) บางกรณียังไม่รวมภาพ Overlay ออกไป\n"
               "แนวทางการแก้ไขในอนาคต: พัฒนา Thumbnail Canvas พรีวิวสดสำหรับ Asset รูปภาพและเสียงใน Properties Panel และปรับปรุงระบบ Video Exporter Composite Graph ใน video_exporter.py ให้รวมไฟล์รูปภาพและแทร็กเสียงอิสระเข้าสู่ไฟล์ส่งออกขั้นสุดท้ายอย่างสมบูรณ์")

    add_bullet("2. การเพิ่มประสิทธิภาพอัตราเฟรมพรีวิววิดีโอระดับสูง (High FPS GPU Hardware Pipeline > 60 FPS)",
               "ปัญหาที่พบ: แม้การแสดงผลพรีวิววิดีโอจะเพิ่มความลื่นไหลเป็น 30-40 FPS แล้ว แต่การวาดภาพบน Tkinter Canvas ยังใช้การแปลงภาพเป็น PhotoImage บน CPU เป็นหลัก ทำให้ยังคงมีเฟรมตกเล็กน้อยบนไฟล์ความละเอียดสูง 4K\n"
               "แนวทางการแก้ไขในอนาคต: พัฒนา Zero-Copy Hardware Canvas Pipeline โดยใช้ PyOpenGL หรือ Direct3D Canvas ร่วมกับ GPU Texture Sharing ในการวาดแสดงผลเฟรมภาพสดแทน Tkinter Canvas เพื่อดันอัตราเฟรมพรีวิววิดีโอขณะตัดต่อให้ลื่นไหลสูงสุดทะลุ 60+ FPS บนการ์ดจอ RTX 3060")

    # Save document
    out_dir = r"D:\Folder_For_Work\Year4_1\Media Pro Project"
    out_path = os.path.join(out_dir, "รายงานสรุปการปรับปรุงแก้ไขและพัฒนาโปรแกรม_MediaPro_Final.docx")
    try:
        doc.save(out_path)
        print(f"REPORT GENERATED SUCCESSFULLY AT: {out_path}")
    except PermissionError:
        out_path_v2 = os.path.join(out_dir, "รายงานสรุปการปรับปรุงแก้ไขและพัฒนาโปรแกรม_MediaPro_Updated.docx")
        doc.save(out_path_v2)
        print(f"REPORT GENERATED SUCCESSFULLY AT FALLBACK PATH: {out_path_v2}")

if __name__ == "__main__":
    create_report()
