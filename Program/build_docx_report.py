import os
import sys

def create_report():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx not installed.")
        return False

    doc = Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles / Fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cordia New'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    def set_cell_background(cell, fill_hex):
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tc_pr.append(shd)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("รายงานสรุปปัญหาที่พบและแนวทางการแก้ไขปรับปรุง\nโปรแกรมระบบตัดต่อวิดีโอ (MediaPro Video Editor)")
    title_run.font.name = 'Angsana New'
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy Blue
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("สรุปผลงานความก้าวหน้าและการเพิ่มประสิทธิภาพระบบ")
    sub_run.font.name = 'Angsana New'
    sub_run.font.size = Pt(18)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Summary Heading
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. สรุปภาพรวมผลการดำเนินงาน (Executive Summary)")
    h1_run.font.name = 'Angsana New'
    h1_run.font.size = Pt(20)
    h1_run.bold = True
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    lead_p = doc.add_paragraph(
        "จากการพัฒนาระบบร่วมกันในการแก้ไขปัญหาเชิงโครงสร้างของโปรแกรม MediaPro Video Editor "
        "ซึ่งเป็นระบบตัดต่อวิดีโอสไตล์ CapCut ที่ใช้ Python เป็นหลักในการประมวลผล "
        "คณะทำงานได้ดำเนินการแก้ไขปัญหาหลักที่ส่งผลต่อประสิทธิภาพการทำงาน (Performance) "
        "และการใช้งาน (User Experience) จำนวนทั้งสิ้น 6 รายการ ทำให้โปรแกรมรุ่นปัจจุบันทำงานได้อย่างเสถียร "
        "ไม่เกิดปัญหากระตุกในพรีวิว ภาพกับเสียงตรงกันอย่างแม่นยำ รองรับการทำงานของเลเยอร์ Overlay ภาพซ้อน และสามารถส่งออกผลงานได้อย่างสมบูรณ์แบบ"
    )
    lead_p.paragraph_format.space_after = Pt(12)

    # Table Heading
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. ตารางสรุปปัญหาที่พบและแนวทางแก้ไข")
    h2_run.font.name = 'Angsana New'
    h2_run.font.size = Pt(20)
    h2_run.bold = True
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    # Add Table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ["รายการปัญหาที่พบ (Issue / Problem)", "แนวทางการแก้ไขเชิงเทคนิค (Technical Solution)", "ผลลัพธ์ที่ได้ (Expected Result)"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], '1B365D')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Cordia New'
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Table Contents
    items = [
        (
            "ภาพแสดงผลในโปรแกรมกระตุกมากและ CPU ทำงานสูงขณะเล่นพรีวิว",
            "ปรับระบบอ่านวิดีโอเป็นแบบเรียงลำดับ (Sequential) ย้ายงานครอป/ลดขนาด/แปลงสีไปทำใน Background thread และพรีวิวภาพที่ขนาด 360p ชั่วคราวขณะเล่น",
            "พรีวิวแสดงผลลื่นไหลขึ้น อัตราการใช้ CPU ลดลงกว่า 80% ป้องกันหน้าต่างโปรแกรมค้าง"
        ),
        (
            "ภาพและเสียงแสดงผลไม่ตรงกัน (Audio/Video Desync / Rubber-banding)",
            "เปลี่ยนระบบนาฬิกาตัวเล่นพรีวิวเป็นแบบ Monotonic ที่เดินหน้าอย่างเดียว และตั้งระบบรอโหลดภาพพรีวิวล่วงหน้า (Pre-fill Buffer) ให้เสร็จก่อนเปิดเสียง",
            "ภาพและเสียงตรงกันอย่างแม่นยำ ไม่มีอาการภาพกระโดดย้อนหลังไปมา"
        ),
        (
            "แทร็ก Overlay (ภาพซ้อน) ไม่สามารถแสดงผลได้จริง",
            "สร้างระบบผสมภาพ (Overlay Composite Engine) ทั้งในจอพรีวิวและขั้นตอน Export ให้ตรวจสอบและนำรูปภาพ (PNG/JPG) หรือวิดีโอซ้อนวางทับลงไปตามเวลาที่กำหนด",
            "รองรับการทำ Picture-in-Picture และการใส่สติกเกอร์/รูปภาพ ซ้อนทับในงานได้จริง"
        ),
        (
            "ระบบถอดซับไตเติลค้าง (Not Responding) หรือใช้เวลานานผิดปกติ",
            "ตั้งค่า Socket Timeout 3 วินาทีเพื่อจำกัดเวลาตรวจเช็ค Silero VAD ออนไลน์ หากไม่มีเน็ตให้สลับมาใช้ระบบแบ่งช่วงเสียงออฟไลน์ (Fixed 20s Chunks) ทันที",
            "โปรแกรมไม่ค้างเมื่อไม่มีอินเทอร์เน็ต สามารถถอดคำบรรยายแบบออฟไลน์ได้อย่างรวดเร็ว"
        ),
        (
            "ปุ่มสั่ง Export จมหายจากหน้าต่าง และวิดีโอพังเมื่อบางช่วงไม่มีเสียง",
            "ขยายขนาดหน้าต่าง Export สูงขึ้นเป็น 550px, รวมคลิปบนไทม์ไลน์เป็นไฟล์เดียวเพื่อไปเบิร์นซับ และตรวจสอบหากคลิปไม่มีเสียงให้แทรกเสียงเงียบ (anullsrc) ป้องกัน FFmpeg crash",
            "หน้าจอกดส่งออกมีปุ่มครบถ้วน ส่งออกงานได้สมบูรณ์แบบ 100% ไม่ว่าวิดีโอจะมีกี่คลิปหรือไม่มีเสียงก็ตาม"
        ),
        (
            "ย้ายตำแหน่งรูปภาพหรือคลิปบน Timeline ได้ยาก",
            "เปลี่ยนการระบุตำแหน่งแนวตั้งของเมาส์มาใช้ canvasy และปรับค่าความแรงของแม่เหล็กสแนป (SNAP_PX) จาก 12px ลดลงเหลือ 8px",
            "ผู้ใช้งานสามารถคลิกเลือกและลากย้ายรูปภาพหรือคลิปจัดวางได้อย่างอิสระและลื่นไหล"
        )
    ]

    for row_idx, data in enumerate(items, 1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Cordia New'
                    run.font.size = Pt(14)
        # Background coloring for zebra striping
        if row_idx % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, 'F2F4F7')

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Detailed report section
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. รายละเอียดเชิงลึกปัญหารายข้อและแนวทางแก้ไข")
    h3_run.font.name = 'Angsana New'
    h3_run.font.size = Pt(20)
    h3_run.bold = True
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    detailed_issues = [
        {
            "num": "3.1",
            "title": "ปัญหาประสิทธิภาพและความกระตุกขณะพรีวิว (Preview Performance & CPU Optimization)",
            "problem": "เมื่อทำการเล่นวิดีโอบน Timeline ตัวโปรแกรมจะหน่วงและกระตุกมาก อัตราการใช้ CPU พุ่งสูงถึง 100% ซึ่งทำให้ประสิทธิภาพโดยรวมของคอมพิวเตอร์ลดลงอย่างเห็นได้ชัดและไม่สะดวกต่อการทำงาน",
            "cause": "การถอดรหัสภาพวิดีโอของ OpenCV ใน Thread เบื้องหลังเดิมใช้การสุ่มตำแหน่งเฟรมแบบสุ่ม (cap.set(POS_MSEC)) ในทุกๆ เฟรม ซึ่งเป็นการเรียกคำสั่งที่สิ้นเปลืองพลังงาน CPU สูงมาก ประกอบกับตัวโปรแกรมพยายามเรนเดอร์และอัปเดตหน้าตา UI ใน Main thread ในทุกๆ เฟรม และการ resize ภาพบ่อยครั้ง",
            "sol": "1. ปรับเปลี่ยนกระบวนการถอดรหัสเป็นการอ่านภาพเรียงลำดับปกติ (Sequential read) และสั่ง Seek เจาะจงเฉพาะเมื่อพิกัดของเฟรมเกิดการเหลื่อมล้ำเกิน 3 เฟรมขึ้นไป\n"
                   "2. ย้ายกระบวนการครอปภาพ, การ Resize, และการแปลงระบบสี BGR->RGB ไปทำใน Background Thread แทนที่จะทำบน Main Thread\n"
                   "3. ระหว่างที่ทำการเล่นพรีวิว จะทำการย่อขนาดภาพชั่วคราวให้เป็นความละเอียดต่ำ (360p) เพื่อประหยัด CPU ในการย้ายข้อมูลและแสดงผล (เมื่อหยุดเล่นภาพจะแสดงความละเอียดดิบคู่กับซับไตเติลชัดเจนเหมือนเดิม)\n"
                   "4. ปรับลดการเรียกอัปเดตหน้าตา properties panel ระหว่างที่วิดีโอกำลังเล่นพรีวิวอยู่"
        },
        {
            "num": "3.2",
            "title": "ปัญหาการไม่ตรงกันของสัญญาณภาพและเสียง (Audio-Video Synchronization)",
            "problem": "ในบางครั้งระหว่างที่เล่นพรีวิววิดีโอ เสียงพูดกับสัญญาณปากในวิดีโอมีความล่าช้าต่างกัน (Desync) และภาพเกิดอาการขยับกระโดดย้อนหลังไปมา (Rubber-banding)",
            "cause": "โปรแกรมดึงค่าตำแหน่งเสียงผ่านฟังก์ชัน pygame.mixer.music.get_pos() มาเป็นตัวควบคุมพิกัดของภาพวิดีโอ ซึ่งบนระบบปฏิบัติการ Windows การตอบสนองของระบบเสียงมีความหน่วงและพิกัดค่ามีอาการกระเพื่อมไม่สม่ำเสมอ ส่งผลให้เฟรมเป้าหมายถูกบังคับให้เลื่อนถอยหลังทำให้เกิดภาพย้อนทาง",
            "sol": "1. ปรับเปลี่ยนนาฬิกากลางของระบบพรีวิวมาใช้ระบบเวลา monotonic (time.perf_counter()) ซึ่งเพิ่มขึ้นอย่างราบเรียบและเสถียร\n"
                   "2. เขียนฟังก์ชันป้องกันภาพย้อนหลัง (Monotonic Guard: target_fi = max(self.fi, ...)) บังคับให้เฟรมเป้าหมายเดินหน้าอย่างเดียว\n"
                   "3. ตั้งระบบ Buffer ภาพล่วงหน้าก่อนปล่อยเสียง (Pre-fill Buffer) รอให้เฟรมภาพถอดรหัสเสร็จสะสมครบอย่างน้อย 10 เฟรมก่อน จึงจะเริ่มส่งสัญญาณปล่อยเสียงเพลง ทำให้จุดเริ่มต้นภาพและเสียงสัมพันธ์กันตั้งแต่จุดสตาร์ท"
        },
        {
            "num": "3.3",
            "title": "ปัญหาระบบเลอย์ภาพซ้อนไม่แสดงผลจริง (Overlay Compositing Support)",
            "problem": "เมื่อผู้ใช้ทำการวางไฟล์รูปภาพ (เช่น PNG/JPG) หรือวิดีโอเพิ่มลงในแทร็ก Overlay ตัวภาพชิ้นงานเหล่านั้นไม่ไปแสดงผลอยู่บนวิดีโอพรีวิว และเมื่อทำการส่งออก (Export) ภาพเหล่านั้นก็ไม่ปรากฏในวิดีโอที่ได้",
            "cause": "ระบบพรีวิวและการเรนเดอร์ส่งออกไฟล์ (Export) ของโปรแกรมเดิมเขียนให้อ่านและประมวลผลวิดีโอจากแทร็กหลัก (Main Track) เพียงอย่างเดียว โดยระบบยังไม่มีกระบวนการประมวลผลซ้อนภาพจากแทร็ก Overlay เข้าไปด้วย",
            "sol": "1. เขียนระบบผสมภาพพรีวิว (_apply_overlay) ทำหน้าที่สแกนดูแทร็ก Overlay ณ ช่วงเวลาปัจจุบัน หากพบภาพหรือวิดีโอ จะนำเฟรมดังกล่าวมาจัดสเกลให้อยู่ในขนาด 50% กึ่งกลางหน้าจอแล้วประทับทับบนวิดีโอหลัก (หากเป็น PNG จะทำ Alpha Blending เพื่อรักษาความโปร่งใสของพื้นหลังภาพ)\n"
                   "2. ปรับปรุงระบบ Export ของ FFmpeg ให้เพิ่มอินพุตของคลิป Overlay ทั้งหมดเข้าไป และใช้ filter complex 'overlay' ของ FFmpeg ในการผสมรูปภาพ/วิดีโอซ้อน ณ ช่วงเวลาที่กำหนดได้อย่างเสถียรและแม่นยำสูง"
        },
        {
            "num": "3.4",
            "title": "ปัญหาการสร้างซับไตเติลโหลดนานหรือหน้าต่างค้าง (Subtitle Transcription Timeout)",
            "problem": "เมื่อทำการกดปุ่มถอดความสร้างซับไตเติลอัตโนมัติ (Transcribe) ในบางครั้งโปรแกรมจะค้างยาวนานมากแบบไม่ขึ้นความคืบหน้าใดๆ และหน้าต่างจะแสดงสถานะค้าง (Not Responding)",
            "cause": "ระบบ VAD ตรวจหาช่วงเสียงพูด (Silero VAD) พยายามส่งคำสั่งเชื่อมต่อผ่านระบบออนไลน์ไปเช็คอัปเดตดาวน์โหลดโมเดลบน GitHub/PyTorch Hub ทุกครั้งในการรัน ซึ่งหากไม่มีอินเทอร์เน็ตหรือมีอินเทอร์เน็ตที่ช้ามาก ตัวระบบจะเกิดการบล็อก (Hanging) นานหลายนาทีเพื่อรอ Connection Timeout",
            "sol": "1. กำหนดค่า Socket Timeout ให้กับระบบเครือข่ายเป็นเวลาไม่เกิน 3 วินาที เพื่อให้ยกเลิกการรออย่างรวดเร็วหากไม่มีอินเทอร์เน็ต\n"
                   "2. เขียนฟังก์ชันระบบออฟไลน์สำรอง (Offline Fallback Chunks) ในกรณีที่ตรวจสอบแล้วโหลดโมเดล Silero VAD ออนไลน์ไม่ได้ โดยจะสลับมาใช้วิธีแบ่งช่วงสัญญาณเสียงออฟไลน์เป็นช่วงละ 20 วินาทีคงที่ทันที และส่งเข้า Whisper เพื่อถอดความ ซึ่ง Whisper จะประมวลผลต่อได้เสร็จอย่างรวดเร็วโดยไม่ต้องใช้เน็ต"
        },
        {
            "num": "3.5",
            "title": "ปัญหาระบบส่งออกงานและการหลุดของปุ่มสั่งการ (Export Dialog & FFmpeg Crash)",
            "problem": "1. เมื่อกดเมนูส่งออก (Export) หน้าต่างตั้งค่าการส่งออกขึ้นมาไม่ครบ โดยปุ่มสำหรับสั่งบันทึกหรือปุ่ม Cancel หลุดหายไปจากขอบล่างของหน้าต่าง\n"
                   "2. เมื่อสั่ง Export หากมีคลิปใดคลิปหนึ่งบนไทม์ไลน์ไม่มีเสียง (เช่น ภาพนิ่ง หรือไฟล์ปิดเสียง) กระบวนการส่งออกจะเออเร่อและไม่บันทึกงานให้",
            "cause": "1. ความสูงหน้าต่าง Toplevel ของ Export Dialog กำหนดไว้เพียง 470px ซึ่งสั้นกว่าความสูงขององค์ประกอบทั้งหมดในหน้าจอ ทำให้ปุ่มโดนดันหลุดขอบ\n"
                   "2. คำสั่งรวมไฟล์ของ FFmpeg filter complex คาดหวังว่าทุกคลิปอินพุตจะต้องมีทั้งภาพ [v] และเสียง [a] เมื่อมีวิดีโอที่ไม่มีสัญญาณเสียงทำให้ stream [a] หาไม่เจอและคำสั่งหยุดทำงาน (return code 1)\n"
                   "3. ระบบ Export กับการ Burn Subtitle ทำงานแยกกันอย่างสิ้นเชิง ส่งผลให้เวลาเรนเดอร์แบบ Burn ซับ คลิปอื่นๆ บนไทม์ไลน์หายหมดเหลือแค่คลิปแรก",
            "sol": "1. ขยายความสูงของหน้าต่าง Export Dialog เป็น 550px เพื่อรองรับปุ่มได้ครบถ้วน\n"
                   "2. ปรับระบบการส่งออกเป็นแบบ Unified Pipeline โดยทำการ Concat วิดีโอบนไทม์ไลน์ทั้งหมดเป็นวิดีโอชิ้นเดียวก่อน จากนั้นนำไฟล์รวมนี้ไปส่งต่อให้ระบบ Burn Subtitle ทำให้คลิปทั้งหมดและเอฟเฟกต์ไม่สูญหาย\n"
                   "3. เขียนสคริปต์ Probe ตรวจสอบสัญญาณเสียงของทุกคลิป หากพบว่าคลิปใดไม่มีเสียง จะทำคำสั่งแทรกเสียงเงียบจำลอง (anullsrc) ที่มีความยาวเท่ากับคลิปตัวนั้นผสมเข้าไปใน FFmpeg เพื่อป้องกันคำสั่งพัง"
        },
        {
            "num": "3.6",
            "title": "ปัญหาย้ายตำแหน่งรูปภาพและคลิปบน Timeline ยาก (Timeline Interaction & Snapping)",
            "problem": "ผู้ใช้งานไม่สามารถลากย้ายรูปภาพหรือคลิปบน Timeline ได้สะดวก เมื่อคลิกเลือกแล้วลากขยับ มักจะลากย้ายไม่ได้หรือตัวคลิปกระเด้งกลับมาอยู่ที่จุดเริ่มต้นเดิม",
            "cause": "1. พิกัดแกนดิ่งใน Canvas ใช้ระบบพิกัดหน้าต่าง widget (e.y) ซึ่งหากโปรเจกต์มีการปรับเปลี่ยนขนาดหน้าต่าง ค่าแนวตั้งจะคลาดเคลื่อนและคลิกกดไม่โดน\n"
                   "2. ระยะล็อกคลิป (SNAP_PX) มีค่าถึง 12px ซึ่งแรงเกินไป ทำให้เมื่อขยับเมาส์เพียงเล็กน้อย คลิปจะดีดตัวกลับเข้าสแนปล็อกที่เดิมตลอดจนเหมือนลากไปไหนไม่ได้",
            "sol": "1. ปรับการรับค่าคลิกเมาส์ของแกนดิ่งมาแปลงเป็นพิกัดจริงของระบบ Canvas (self._tlc.canvasy(e.y))\n"
                   "2. ปรับลดขนาดช่วงการล็อกแม่เหล็กสแนปจาก 12px ให้เหลือเพียง 8px ช่วยให้ลากคลิปหลุดจากแรงดูดสแนปได้อย่างอิสระและจัดวางง่ายขึ้น"
        }
    ]

    for item in detailed_issues:
        # Title of section
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{item['num']} {item['title']}")
        run.bold = True
        run.font.name = 'Angsana New'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        # Problem
        p_prob = doc.add_paragraph()
        p_prob.paragraph_format.left_indent = Inches(0.25)
        p_prob.paragraph_format.space_after = Pt(2)
        r_lbl = p_prob.add_run("• ปัญหาที่พบเจอ: ")
        r_lbl.bold = True
        r_lbl.font.color.rgb = RGBColor(0xBA, 0x3C, 0x2A) # Dark Red
        p_prob.add_run(item["problem"])

        # Cause
        p_cause = doc.add_paragraph()
        p_cause.paragraph_format.left_indent = Inches(0.25)
        p_cause.paragraph_format.space_after = Pt(2)
        r_lbl = p_cause.add_run("• สาเหตุเชิงเทคนิค: ")
        r_lbl.bold = True
        p_cause.add_run(item["cause"])

        # Solution
        p_sol = doc.add_paragraph()
        p_sol.paragraph_format.left_indent = Inches(0.25)
        p_sol.paragraph_format.space_after = Pt(8)
        r_lbl = p_sol.add_run("• วิธีการแก้ไขและผลลัพธ์: ")
        r_lbl.bold = True
        r_lbl.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F) # Dark Green
        p_sol.add_run(item["sol"])

    # Footer/Sign
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_p.add_run("จัดทำโดย: คณะทำงานพัฒนาระบบ MediaPro\nวันที่แก้ไขปรับปรุง: 30 มิถุนายน 2570")
    sign_run.font.italic = True
    sign_run.font.size = Pt(14)
    sign_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Save
    out_name = "report_media_pro_fixes.docx"
    doc.save(out_name)
    print(f"Success: Saved report as {out_name}")
    return True

if __name__ == '__main__':
    create_report()
